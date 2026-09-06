from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81) # Deep blue

def add_bold_bullet(doc, title, text):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(title)
    r1.bold = True
    r2 = p.add_run(text)

def main():
    doc = Document()

    # Title
    title = doc.add_heading('Laporan Analisis Detail Teknis dan Proses Bisnis\nPortal PORPROV XV Jawa Barat 2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    p = doc.add_paragraph('Dokumen ini membedah secara mendalam (langkah demi langkah) bagaimana operasional sistem PORPROV XV Kota Depok berjalan, mulai dari input data dasar hingga publikasi skor secara real-time.')

    # 1
    add_heading(doc, '1. Detail Proses Bisnis & Alur Kerja (Workflow)', 1)
    doc.add_paragraph('Inti dari aplikasi ini adalah memastikan data pertandingan mengalir dengan valid. Berikut adalah alur eksekusi di lapangan:')

    add_heading(doc, 'A. Persiapan Ekosistem (Master Data)', 2)
    doc.add_paragraph('Sebelum pertandingan bisa dijadwalkan, pondasi data wajib didirikan oleh Admin Organisasi / Super Admin:')
    add_bold_bullet(doc, 'Input Cabang Olahraga (Cabor): ', 'Admin masuk ke modul Master Data -> Cabor. Menambahkan nama cabang olahraga (misal: "Renang", "Pencak Silat"). Data ini akan menjadi induk dari semua pertandingan.')
    add_bold_bullet(doc, 'Input Kontingen (Peserta Daerah): ', 'Admin masuk ke modul Master Data -> Kontingen. Mendaftarkan seluruh kontingen perwakilan Kota/Kabupaten peserta PORPROV.')
    add_bold_bullet(doc, 'Input Venue (Lokasi): ', 'Admin masuk ke modul Venue. Memasukkan nama stadion/GOR, kapasitas, dan WAJIB menyertakan titik koordinat (Latitude & Longitude) agar terintegrasi dengan Peta (City Guide).')
    add_bold_bullet(doc, 'Manajemen Atlet/Tim: ', 'Entitas spesifik dari masing-masing Kontingen didaftarkan sebagai kandidat yang akan bertanding di nomor tertentu.')

    add_heading(doc, 'B. Manajemen Jadwal & Nomor Pertandingan', 2)
    doc.add_paragraph('Setelah fondasi siap, Admin Pertandingan akan merakit jadwal dengan cara:')
    doc.add_paragraph('1. Membuka modul Jadwal Pertandingan (Schedule).', style='List Number')
    doc.add_paragraph('2. Memilih Cabang Olahraga yang telah didaftarkan.', style='List Number')
    doc.add_paragraph('3. Memilih Venue tempat bertanding.', style='List Number')
    doc.add_paragraph('4. Mengisi Waktu Mulai (Start Time) pertandingan.', style='List Number')
    doc.add_paragraph('5. Menentukan Peserta (Nomor Pertandingan): Admin menetapkan formasi peserta (Peserta A vs Peserta B). Sistem menjamin keakuratan dengan memaksa tipe peserta seragam—jika A adalah Individu, maka B harus Individu (nama atlet merujuk ke afiliasi Kontingen). Jika beregu, keduanya harus Tim.', style='List Number')

    add_heading(doc, 'C. Pelaksanaan Pertandingan (LiveScore)', 2)
    doc.add_paragraph('Saat peluit dibunyikan, Koresponden (Petugas Lapangan) mengambil alih:')
    doc.add_paragraph('1. Koresponden membuka modul LiveScore Center di Web Admin.', style='List Number')
    doc.add_paragraph('2. Memilih Jadwal Pertandingan yang sedang berlangsung saat itu.', style='List Number')
    doc.add_paragraph('3. Melakukan pembaruan (Update) skor kubu A dan kubu B secara berkala.', style='List Number')
    doc.add_paragraph('4. Sistem di belakang layar (Event-Driven) otomatis memancarkan skor baru ini ke seluruh pengunjung Web Publik secara langsung tanpa tunda (Real-Time).', style='List Number')

    add_heading(doc, 'D. Penentuan Juara (Alur Medali Berjenjang)', 2)
    doc.add_paragraph('Proses pencatatan medali harus bebas dari manipulasi. Tahapannya:')
    add_bold_bullet(doc, 'Tahap Pengajuan (Submit): ', 'Koresponden menekan tombol "Kirim Hasil" di akhir pertandingan. Status menjadi PENDING.')
    add_bold_bullet(doc, 'Tahap Validasi (Verify): ', 'Verifikator (Juri) memeriksa pengajuan tersebut. Jika sah, ditekan "Verify" (status VERIFIED). Jika sengketa, tekan "Reject".')
    add_bold_bullet(doc, 'Tahap Pengesahan (Publish): ', 'Super Admin mengklik "Publish". Barulah medali diakui secara OFFICIAL dan masuk ke Klasemen Umum Publik.')

    # 2
    add_heading(doc, '2. Detail Layanan Utama', 1)
    add_bold_bullet(doc, 'Layanan Publik (Tanpa Login): ', 'Akses anonim bagi warga untuk membuka Homepage, Klasemen, LiveScore, dan City Guide. Dioptimalkan penuh (Next.js SSR & PWA) agar tidak lumpuh meski diakses jutaan warga Jabar berbarengan.')
    add_bold_bullet(doc, 'Layanan Autentikasi (Keycloak): ', 'Pintu masuk SSO (Single Sign-On). Jika ada staf mencoba masuk Web Admin, ia dilempar ke gerbang Keycloak untuk verifikasi Username, Password, dan Role.')
    add_bold_bullet(doc, 'Layanan Server-Sent Events (SSE) Stream: ', 'Terowongan data asinkron. Bukannya warga yang memuat ulang web (Tarik data), layanan ini yang "Mendorong" data skor baru langsung ke layar perangkat warga secara otomatis seperti notifikasi chat.')

    # 3
    add_heading(doc, '3. Cara Kerja Fitur Lengkap (Mekanisme)', 1)
    add_heading(doc, 'A. Fitur Recycle Bin (Soft Delete)', 2)
    doc.add_paragraph('Tidak ada penghapusan data permanen akibat kelalaian.')
    doc.add_paragraph('Cara Kerja: Jika Admin menghapus data Venue, sistem hanya memberikan stempel deleted_at. Data disembunyikan dari daftar utama, namun masih ada di Recycle Bin.', style='List Bullet')
    doc.add_paragraph('Pemulihan: Admin membuka modul Recycle Bin dan menekan "Restore", data kembali sempurna beserta relasinya.', style='List Bullet')

    add_heading(doc, 'B. Papan Penelusuran Mutasi (Audit Log)', 2)
    doc.add_paragraph('Kamera pengawas sistem (Immutable Log) untuk mencegah kecurangan.')
    doc.add_paragraph('Setiap aksi (Tambah Cabor, Hapus Jadwal, Ubah Skor) dicatat permanen oleh sistem Audit.', style='List Bullet')
    doc.add_paragraph('Auditor dapat melihat rekam jejak mendetail ("Pegawai A dari IP sekian, mengubah skor B") dan mengekspornya ke format CSV.', style='List Bullet')

    add_heading(doc, 'C. Manajemen Tampilan Publik (Hero Landing Page)', 2)
    doc.add_paragraph('Tampilan papan promosi beranda Web Publik bersifat dinamis.')
    doc.add_paragraph('Humas dapat mengganti teks promosi dan gambar dari modul Hero Utama di Admin.', style='List Bullet')
    doc.add_paragraph('Gambar harus dipilih langsung dari integrasi Media Library.', style='List Bullet')

    # 4
    add_heading(doc, '4. Aspek Keamanan & Ketahanan Infrastruktur', 1)
    add_bold_bullet(doc, 'Dockerisasi: ', 'Layanan dipecah dalam kontainer independen. Jika modul Venue kelebihan beban, modul LiveScore akan tetap hidup dan beroperasi normal.')
    add_bold_bullet(doc, 'Koneksi Database Konkuren (pgxpool): ', 'Backend Golang menggunakan pool koneksi. Ia dapat memproses ribuan permintaan database secara berbarengan (konkuren) tanpa menyebabkan bottleneck.')
    add_bold_bullet(doc, 'Security Headers (CSP & CORS): ', 'Sistem secara otomatis menolak upaya peretas yang mencoba menginjeksi naskah jahat (XSS). Font dan aset ditanam di server sendiri (self-hosted) agar mandiri dan tertutup.')

    # Conclusion
    doc.add_paragraph("\n")
    p_end = doc.add_paragraph('Laporan ini disusun dengan pendekatan prosedural operasional lapangan, sehingga para operator, manajer IT, hingga end-user dapat merelasikan fitur sistem dengan aktivitas pertandingan di dunia nyata secara gamblang.', style='Quote')
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save doc
    output_path = Path(__file__).resolve().parent / "Laporan_Detail_Operasional_Porprov.docx"
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    main()
