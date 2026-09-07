from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81) # Deep blue

def main():
    doc = Document()

    # Title
    title = doc.add_heading('Laporan Analisis Komprehensif\nPortal PORPROV XV Jawa Barat 2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    p = doc.add_paragraph('Selamat datang di tinjauan menyeluruh Portal PORPROV XV! Dokumen ini dirancang agar mudah dicerna, asyik untuk dibaca, dan sangat pas dijadikan pegangan bagi Anda yang akan mempresentasikannya di depan pimpinan maupun audiens umum.\n\nMari kita bedah kehebatan sistem yang sudah hidup dan bernapas di VPS Publik kita.')

    # 1
    add_heading(doc, '1. Proses Bisnis: Urat Nadi Pertandingan', 1)
    doc.add_paragraph('Aplikasi ini bukan sekadar web pameran, melainkan mesin penggerak roda pertandingan yang sesungguhnya.')
    doc.add_paragraph('Manajemen Terpusat: Semua entitas olahraga (Cabang Olahraga, Venue, Kontingen, hingga Jadwal Pertandingan) dikendalikan dari satu pintu yang rapi.', style='List Bullet')
    doc.add_paragraph('Sirkulasi Medali Berjenjang: Proses klaim medali tidak sembarangan. Dimulai dari Koresponden di lapangan yang menginput data, lalu diverifikasi oleh Verifikator agar terhindar dari kecurangan, hingga akhirnya disahkan (Publish) oleh Super Admin untuk masuk ke Klasemen Umum.', style='List Bullet')
    doc.add_paragraph('Siaran Langsung Skor (LiveScore): Penonton tidak perlu menunggu besok pagi baca koran; skor pertandingan mengalir detik itu juga (Real-Time) langsung dari ujung jari petugas di lapangan ke layar gawai masyarakat.', style='List Bullet')

    # 2
    add_heading(doc, '2. Layanan: Tiga Pilar Utama', 1)
    doc.add_paragraph('Sistem memecah pelayanannya menjadi tiga ranah yang fokus pada target penggunanya:')
    doc.add_paragraph('Web Publik (Public Portal): Wajah depan PORPROV untuk masyarakat umum. Tanpa perlu login, siapa saja bisa melihat klasemen, jadwal, skor terkini, dan informasi pariwisata (City Guide) Kota Depok.', style='List Bullet')
    doc.add_paragraph('Web Admin (Back-Office): Dapur rahasia para operator dan panitia. Dilengkapi keamanan tinggi, sistem ini melayani manajemen data, moderasi, audit, dan kontrol akses berjenjang.', style='List Bullet')
    doc.add_paragraph('Gateway Waktu Nyata (Real-Time Stream): Layanan tak kasat mata yang terus-menerus memancarkan pembaruan data secara langsung (Server-Sent Events) sehingga layar publik bisa berkedip menampilkan skor baru tanpa perlu di-refresh.', style='List Bullet')

    # 3
    add_heading(doc, '3. Arsitektur Aplikasi: Modern & Tahan Banting', 1)
    doc.add_paragraph('Sistem ini membuang gaya lama (monolith) dan beralih ke arsitektur Microservices (Layanan Mikro) yang gesit:')
    doc.add_paragraph('Pecah-Bagi Tugas: Layanan Master Data, Schedule, Venue, LiveScore, Medal, dan Audit beroperasi secara independen. Jika satu layanan sibuk, layanan lain tetap santai.', style='List Bullet')
    doc.add_paragraph('Event-Driven (NATS): Sistem saling mengobrol melalui jalur komunikasi super cepat (NATS JetStream). Misalnya, jika Venue diubah, layanan Schedule akan otomatis tahu tanpa harus saling menunggu.', style='List Bullet')
    doc.add_paragraph('Pintu Gerbang Tunggal (API Gateway): Seluruh lalu lintas data diatur oleh satu gerbang yang canggih. Gateway ini bertugas memeriksa karcis (Token), memastikan izin (RBAC), lalu mengarahkan ke layanan mikro yang tepat.', style='List Bullet')
    doc.add_paragraph('SSO (Single Sign-On) Keycloak: Tidak ada lagi repot bikin banyak akun. Semua otentikasi diurus terpusat oleh Keycloak dengan standar keamanan kelas dunia.', style='List Bullet')

    # 4
    add_heading(doc, '4. Teknologi Aplikasi: Stack Masa Depan', 1)
    doc.add_paragraph('Berdiri di atas fondasi teknologi yang sering dipakai raksasa teknologi modern:')
    doc.add_paragraph('Frontend (Wajah): Next.js & React (TypeScript) dipadukan dengan desain Techwind UI/Tailwind v4. Hasilnya? Tampilan yang memanjakan mata, sangat responsif di layar sentuh, dan disukai mesin pencari (SEO).', style='List Bullet')
    doc.add_paragraph('Backend (Otak): Golang (Go). Bahasa pemrograman super cepat buatan Google yang sangat irit memori dan jago menangani lalu lintas data tinggi secara bersamaan.', style='List Bullet')
    doc.add_paragraph('Database (Ingatan): PostgreSQL sebagai penyimpan data utama, dan Redis sebagai cache (ingatan jangka pendek) agar akses data kilat tanpa membebani server database.', style='List Bullet')
    doc.add_paragraph('Infrastruktur: Seluruh sistem dibungkus dalam Docker Containers dan diorkestrasi via Docker Compose. Dipandu oleh Nginx sebagai pengatur lalu lintas HTTP/HTTPS dengan gembok pengaman SSL.', style='List Bullet')

    # 5
    add_heading(doc, '5. Fitur Lengkap: Senjata Tempur PORPROV', 1)
    doc.add_paragraph('LiveScore Center: Papan skor interaktif yang berdenyut selaras dengan tempo di lapangan.', style='List Bullet')
    doc.add_paragraph('Klasemen Medali (Medal Standing): Klasemen otomatis yang kebal manipulasi karena harus melewati tiga lapis persetujuan.', style='List Bullet')
    doc.add_paragraph('Master Data Ekosistem: Pengelolaan Atlet, Tim, Kontingen, dan Cabang Olahraga terintegrasi.', style='List Bullet')
    doc.add_paragraph('Jejak Audit (Immutable Audit Log): "Kamera pengawas" sistem. Setiap perubahan sekecil apa pun dicatat permanen; ketahuan siapa yang mengubah, kapan, dan apa yang diubah.', style='List Bullet')
    doc.add_paragraph('Manajemen Destinasi (City Guide): Fitur bonus untuk mendongkrak pariwisata Depok, lengkap dengan integrasi peta presisi tinggi.', style='List Bullet')

    # 6
    add_heading(doc, '6. Rekomendasi Fitur Tambahan (Next-Gen)', 1)
    doc.add_paragraph('Untuk membawa aplikasi ini ke level selanjutnya, berikut ide brilian yang bisa dipertimbangkan:')
    doc.add_paragraph('Aplikasi Mobile Spesifik (React Native): Walau versi web sudah Mobile-First, merilis versi Android/iOS akan memberikan prestise khusus dan memungkinkan penggunaan fitur kamera (misal: scan tiket/QR atlet).', style='List Bullet')
    doc.add_paragraph('Push Notification & Integrasi WhatsApp: Fitur yang mengirimi notifikasi ke panitia jika ada pengajuan medali baru, atau peringatan jika jadwal mendadak berubah.', style='List Bullet')
    doc.add_paragraph('Peta Venue Interaktif (Live Mapping): Pengunjung publik dapat melihat titik sebaran venue di Kota Depok pada satu peta digital interaktif beserta tingkat kepadatan atau kondisi cuaca di lokasi tersebut.', style='List Bullet')
    doc.add_paragraph('Kemitraan/Sponsor Corner: Spot iklan dinamis pada Web Publik yang perputarannya dikontrol dari Web Admin untuk menggenjot pendapatan daerah.', style='List Bullet')

    # 7
    add_heading(doc, '7. Kebutuhan Sistem & Infrastruktur', 1)
    doc.add_paragraph('Apa yang membuat monster ini tetap menyala?')
    doc.add_paragraph('Lingkungan Saat ini: VPS Publik (192.168.19.3) Linux Ubuntu dengan RAM yang memadai (idealnya 8GB - 16GB) untuk menjalankan puluhan container Docker sekaligus.', style='List Bullet')
    doc.add_paragraph('Keamanan Jaringan: Hanya port Nginx (80/443) yang dibuka ke internet. Komunikasi internal antar layanan mikro sepenuhnya tersembunyi dari publik.', style='List Bullet')
    doc.add_paragraph('Pengamanan SSL: Penggunaan Wildcard SSL (depok.go.id) agar data yang wara-wiri di udara terenkripsi dan terhindar dari peretasan (Gembok Hijau).', style='List Bullet')

    # 8
    add_heading(doc, '8. Aspek Lainnya: Aman Tanpa Kompromi', 1)
    doc.add_paragraph('Soft-Delete Mutlak: Data yang "dihapus" tidak pernah benar-benar hilang dari database, hanya disembunyikan. Ini mencegah tragedi hilangnya data penting akibat salah klik (fitur Recycle Bin).', style='List Bullet')
    doc.add_paragraph('Pembatasan Lintas Batas (CORS & CSP): Peramban pengunjung diawasi ketat. Sistem secara cerdas akan memblokir segala bentuk injeksi kode jahat dari situs tak diundang.', style='List Bullet')
    doc.add_paragraph('Kenyamanan (Accessibility): Desain warna dan kontras mematuhi standar internasional (WCAG 2.2) sehingga tetap nyaman dibaca di bawah terik matahari lapangan maupun mode gelap (Dark Mode) saat malam hari.', style='List Bullet')

    # Conclusion
    doc.add_paragraph("\n")
    p_end = doc.add_paragraph('Aplikasi PORPROV XV ini bukan sekadar pelengkap administratif, melainkan panggung digital berkelas enterprise yang merepresentasikan kemajuan teknologi dan kesiapan Kota Depok sebagai tuan rumah yang inovatif!', style='Quote')
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save doc
    output_path = Path(__file__).resolve().parent / "Laporan_Analisis_Komprehensif_Porprov.docx"
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    main()
