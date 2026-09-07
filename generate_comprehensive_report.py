from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(s)

def add_colored_heading(doc, text, level):
    h = doc.add_heading(text, level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

def bb(doc, title, text):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(title)
    r1.bold = True
    p.add_run(text)

def nb(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        set_cell_shading(c, '0F4C81')
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10)
    style.font.name = 'Calibri'

    # ===== COVER =====
    for _ in range(6): doc.add_paragraph('')
    t = doc.add_heading('LAPORAN ANALISIS KOMPREHENSIF', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2 = doc.add_heading('Portal PORPROV XV Jawa Barat 2026\nKota Depok', 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    p = doc.add_paragraph('Subdomain Publik: https://porprov.depok.go.id')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph('Dinas Komunikasi dan Informatika Kota Depok\nAgustus 2026')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ===== DAFTAR ISI =====
    add_colored_heading(doc, 'Daftar Isi', 1)
    toc = [
        '1. Ringkasan Eksekutif',
        '2. Proses Bisnis Lengkap (Alur Kerja Lapangan)',
        '3. Arsitektur Sistem & Layanan Mikro',
        '4. Teknologi yang Digunakan',
        '5. Web Publik (Public Portal) — Fitur Detail',
        '6. Web Admin (Back-Office) — Fitur Detail',
        '7. Backend Microservices — Detail Teknis',
        '8. Infrastruktur & Keamanan',
        '9. Pemetaan Hak Akses (RBAC)',
        '10. Rekomendasi Fitur Tambahan',
        '11. Kebutuhan Sistem',
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # ===== 1. RINGKASAN EKSEKUTIF =====
    add_colored_heading(doc, '1. Ringkasan Eksekutif', 1)
    doc.add_paragraph(
        'Portal PORPROV XV Jawa Barat 2026 adalah sistem informasi pertandingan olahraga '
        'tingkat provinsi berbasis web yang dibangun dengan arsitektur Microservices modern. '
        'Sistem ini melayani dua kelompok pengguna utama: Masyarakat Umum (melalui Web Publik) '
        'dan Panitia Penyelenggara (melalui Web Admin). Seluruh komponen berjalan di atas '
        'infrastruktur Docker pada VPS Publik Diskominfo Kota Depok (192.168.19.3) dengan '
        'domain resmi https://porprov.depok.go.id.'
    )
    doc.add_paragraph(
        'Sistem terdiri dari 9 layanan mikro Golang, 2 aplikasi web (Next.js & React), '
        '7 database PostgreSQL terisolasi, cache Redis, message broker NATS JetStream, '
        'identity server Keycloak, reverse proxy Nginx dengan SSL, serta observabilitas '
        'Prometheus & Grafana. Total 26 kontainer Docker beroperasi secara simultan.'
    )
    doc.add_page_break()

    # ===== 2. PROSES BISNIS =====
    add_colored_heading(doc, '2. Proses Bisnis Lengkap (Alur Kerja Lapangan)', 1)

    add_colored_heading(doc, '2.1 Tahap Persiapan: Membangun Pondasi Data', 2)
    doc.add_paragraph('Sebelum satu pertandingan pun bisa dijadwalkan, Admin Organisasi / Super Admin wajib mendirikan pondasi data turnamen melalui Web Admin:')

    add_colored_heading(doc, '2.1.1 Input Cabang Olahraga (Cabor)', 3)
    doc.add_paragraph('Admin membuka menu Master Data > Cabang Olahraga, lalu mengisi formulir:')
    add_table(doc, ['Field Form', 'Tipe Input', 'Keterangan'], [
        ['Nama Cabor', 'Teks (wajib)', 'Contoh: "Sepak Bola", "Pencak Silat"'],
        ['Logo Cabor', 'Pemilih Media Library', 'Gambar dipilih dari galeri Media Library'],
        ['Kategori', 'Dropdown', 'Pilihan: Tanding, Seni/Terukur, E-Sports, Eksibisi'],
        ['Status', 'Dropdown', 'Pilihan: Aktif, Eksibisi, Non-Aktif'],
        ['Total Medali', 'Angka', 'Jumlah medali yang diperebutkan (min: 0)'],
        ['Technical Delegate', 'Teks', 'Nama penanggung jawab teknis cabor'],
        ['Deskripsi', 'Teks Panjang', 'Keterangan tambahan tentang cabor'],
    ])

    add_colored_heading(doc, '2.1.2 Input Nomor Pertandingan', 3)
    doc.add_paragraph('Setiap Cabor memiliki sub-kategori bertanding yang disebut "Nomor Pertandingan". Admin mengisi:')
    add_table(doc, ['Field Form', 'Tipe Input', 'Keterangan'], [
        ['Cabang Olahraga', 'Dropdown Pencarian', 'Memilih dari Cabor yang sudah didaftarkan'],
        ['Nama Nomor', 'Teks (wajib)', 'Contoh: "Tunggal Putra", "Ganda Campuran"'],
        ['Kategori Gender', 'Dropdown', 'Pilihan: Putra, Putri, Campuran, Terbuka'],
        ['Tipe Pertandingan', 'Dropdown', 'Pilihan: Tanding, Seni, Terukur, Beregu'],
    ])

    add_colored_heading(doc, '2.1.3 Input Kontingen (Peserta Daerah)', 3)
    doc.add_paragraph('Admin mendaftarkan setiap Kota/Kabupaten peserta PORPROV:')
    add_table(doc, ['Field Form', 'Tipe Input', 'Keterangan'], [
        ['Nama Kontingen', 'Teks (wajib)', 'Contoh: "Kota Depok", "Kota Bandung"'],
        ['Tipe Daerah', 'Dropdown', 'Pilihan: Kota atau Kabupaten'],
        ['Logo Kontingen', 'Pemilih Media Library', 'Logo resmi daerah (opsional)'],
    ])

    add_colored_heading(doc, '2.1.4 Input Venue (Lokasi Pertandingan)', 3)
    doc.add_paragraph('Admin Venue mendaftarkan setiap GOR/Stadion tempat bertanding:')
    add_table(doc, ['Field Form', 'Tipe Input', 'Keterangan'], [
        ['Nama Venue', 'Teks (wajib)', 'Contoh: "GOR Kartika Kostrad"'],
        ['Status Kesiapan', 'Dropdown', 'Pilihan: Persiapan, Siap, Sedang Digunakan'],
        ['Kapasitas', 'Angka (wajib)', 'Kapasitas penonton (kursi)'],
        ['Pilih Cabor', 'Multi-Select Tag', 'Cabor yang akan bertanding di venue ini'],
        ['Alamat Lengkap', 'Teks Panjang', 'Alamat fisik venue'],
        ['Latitude', 'Desimal', 'Koordinat lintang (-90 s.d. 90), default: -6.4025'],
        ['Longitude', 'Desimal', 'Koordinat bujur (-180 s.d. 180), default: 106.7942'],
        ['URL Google Maps', 'URL HTTPS', 'Link rute resmi Google Maps (opsional)'],
        ['Gambar Venue', 'Pemilih Media Library', 'Foto venue dari galeri'],
        ['Contact Person', 'Teks', 'Nama / Nomor HP PJ venue'],
        ['Fasilitas', 'Teks', 'Contoh: Toilet, Parkir, Ruang Medis'],
    ])

    add_colored_heading(doc, '2.1.5 Input City Guide (Panduan Wisata Kota)', 3)
    doc.add_paragraph('Admin memasukkan destinasi wisata, kuliner, dan fasilitas Kota Depok:')
    add_table(doc, ['Field Form', 'Tipe Input', 'Keterangan'], [
        ['Judul', 'Teks (wajib)', 'Nama destinasi, maks 255 karakter'],
        ['Kategori', 'Dropdown', '8 pilihan: Coffee Shop, Wisata Kuliner, Tempat Menginap, Wisata Buatan, Wisata Situ, Pusat Perbelanjaan, Rumah Sakit, Lainnya'],
        ['Deskripsi', 'Teks Panjang', 'Keterangan tentang destinasi'],
        ['Alamat', 'Teks Panjang', 'Alamat fisik destinasi'],
        ['URL Google Maps', 'URL HTTPS', 'Link rute Google Maps resmi (maks 2048 karakter)'],
        ['Latitude & Longitude', 'Desimal (wajib)', 'Koordinat pasangan wajib untuk integrasi peta'],
        ['Tombol Lokasi Saat Ini', 'Tombol GPS', 'Mengambil koordinat otomatis dari perangkat'],
        ['Gambar', 'Pemilih Media Library', 'Foto destinasi dari galeri'],
    ])

    add_colored_heading(doc, '2.2 Tahap Penjadwalan: Merakit Nomor Pertandingan', 2)
    doc.add_paragraph('Setelah pondasi data lengkap, Admin Pertandingan membuat jadwal:')
    add_table(doc, ['Field Form', 'Tipe Input', 'Keterangan'], [
        ['Cabor / Nomor Tanding', 'Dropdown Pencarian (wajib)', 'Memilih gabungan Cabor + Nomor + Gender + Tipe'],
        ['Lokasi (Venue)', 'Dropdown Pencarian (wajib)', 'Memilih dari daftar venue yang sudah terdaftar'],
        ['Waktu Pertandingan', 'Tanggal & Jam (wajib)', 'Format datetime-local'],
        ['Babak', 'Dropdown', 'Pilihan: Penyisihan, Perempat Final, Semifinal, Final'],
        ['Status', 'Dropdown', 'Pilihan: Terjadwal, Berlangsung, Selesai, Ditunda'],
        ['Jenis Peserta', 'Dropdown (seragam)', 'Harus sama untuk A & B: Kontingen, Individu, atau Tim'],
        ['Peserta A — Kontingen', 'Dropdown Pencarian (wajib)', 'Afiliasi daerah slot 1'],
        ['Peserta A — Nama Atlet', 'Teks', 'Wajib jika jenis = Individu (maks 100 karakter)'],
        ['Peserta A — Nama Tim', 'Teks', 'Wajib jika jenis = Tim (maks 150 karakter)'],
        ['Peserta B — Kontingen', 'Dropdown Pencarian (wajib)', 'Afiliasi daerah slot 2'],
        ['Peserta B — Nama Atlet', 'Teks', 'Wajib jika jenis = Individu'],
        ['Peserta B — Nama Tim', 'Teks', 'Wajib jika jenis = Tim'],
    ])
    doc.add_paragraph('Validasi penting: Peserta A dan B wajib bertipe sama dan tidak boleh identik.')

    add_colored_heading(doc, '2.3 Tahap Pelaksanaan: LiveScore Real-Time', 2)
    doc.add_paragraph('Saat hari-H pertandingan, Koresponden (petugas lapangan) melakukan:')
    doc.add_paragraph('1. Login ke Web Admin dengan akun berperan "koresponden".', style='List Number')
    doc.add_paragraph('2. Membuka menu LiveScore Center.', style='List Number')
    doc.add_paragraph('3. Memilih pertandingan yang sedang berlangsung dari dropdown.', style='List Number')
    doc.add_paragraph('4. Mengisi skor Peserta A dan Peserta B (angka non-negatif).', style='List Number')
    doc.add_paragraph('5. Memilih status laga: Belum Mulai, Berlangsung, Istirahat, Selesai, atau Official.', style='List Number')
    doc.add_paragraph('6. Menekan tombol Simpan. Skor langsung terpancar ke seluruh layar publik secara real-time.', style='List Number')
    doc.add_paragraph('')
    doc.add_paragraph('Fitur Koreksi Skor: Jika terjadi kesalahan input, centang "Ini adalah koreksi skor", isi alasan (minimal 5 karakter), lalu simpan. Riwayat koreksi tercatat permanen (append-only, tidak bisa dihapus).')

    add_colored_heading(doc, '2.4 Tahap Penentuan Juara: Workflow Medali 4 Tahap', 2)
    doc.add_paragraph('Proses pencatatan medali melewati 4 tahapan ketat anti-manipulasi:')
    add_table(doc, ['Tahap', 'Status', 'Pelaku', 'Aksi'], [
        ['1. Pengajuan', 'PENDING', 'Koresponden', 'Mengisi kontingen penerima, jumlah emas/perak/perunggu, URL bukti, dan catatan. Menekan "Ajukan Medali".'],
        ['2. Verifikasi', 'VERIFIED', 'Verifikator', 'Memeriksa pengajuan. Menekan "Verifikasi" jika sah, atau "Tolak" dengan alasan (min 5 karakter).'],
        ['3. Pengesahan', 'OFFICIAL', 'Super Admin', 'Menekan "Publikasikan". Medali resmi masuk ke Klasemen Umum Publik.'],
        ['4. Penolakan', 'REJECTED', 'Verifikator/Admin', 'Dapat dilakukan dari status PENDING atau VERIFIED disertai alasan wajib.'],
    ])
    doc.add_paragraph('Setiap aktor (submitted_by, verified_by, rejected_by, published_by) dicatat terpisah dan permanen.')

    add_colored_heading(doc, '2.5 Tahap Editorial: Manajemen Hero Landing Page', 2)
    doc.add_paragraph('Bagian Humas/Pemasaran dapat mengubah tampilan depan Web Publik:')
    add_table(doc, ['Field Form', 'Tipe Input', 'Keterangan'], [
        ['Judul Hero', 'Teks (wajib)', 'Maks 180 karakter. Contoh: "Panggung Juara Jawa Barat."'],
        ['Teks Sorotan', 'Teks', 'Kata kunci yang akan diwarnai gradient (maks 100 karakter, harus substring dari Judul)'],
        ['Isi Hero', 'Teks Panjang (wajib)', 'Deskripsi promosi (maks 1200 karakter)'],
        ['Gambar Latar', 'Pemilih Media Library (wajib)', 'Foto panoramik untuk latar belakang hero'],
        ['Status Aktif', 'Checkbox', 'Hanya 1 Hero yang boleh aktif secara bersamaan'],
    ])
    doc.add_page_break()

    # ===== 3. ARSITEKTUR =====
    add_colored_heading(doc, '3. Arsitektur Sistem & Layanan Mikro', 1)
    doc.add_paragraph('Sistem mengadopsi arsitektur Microservices + Event-Driven yang terdiri dari 9 layanan mikro independen:')
    add_table(doc, ['No', 'Layanan Mikro', 'Port Internal', 'Database', 'Fungsi Utama'], [
        ['1', 'API Gateway', '8000', '— (tanpa DB)', 'Gerbang tunggal: validasi JWT, RBAC, CORS, routing, error masking'],
        ['2', 'Master Data Service', '8081', 'master_data_db', 'Cabor, Kontingen, Nomor Tanding, City Guide, Media Library, Hero'],
        ['3', 'Schedule Service', '8082', 'schedule_db', 'Jadwal pertandingan, susunan peserta A/B, enrichment data'],
        ['4', 'Venue Service', '8087', 'venue_db', 'Venue/GOR, koordinat, kapasitas, fasilitas, relasi cabor/city guide'],
        ['5', 'LiveScore Service', '8083', 'livescore_db', 'Skor real-time, revisi append-only, koreksi, transactional outbox'],
        ['6', 'Medal Standing Service', '8086', 'porprov_db', 'Workflow medali 4 tahap, klasemen resmi, outbox event'],
        ['7', 'Audit Service', '8084', 'audit_db', 'Log mutasi immutable (WORM), SHA-256 hash, subscriber NATS'],
        ['8', 'Realtime Gateway', '8085', 'Redis (cache)', 'SSE stream publik & admin, replay cache, heartbeat, rate limit'],
        ['9', 'User Service', '8001', 'user_service_db', 'Sinkronisasi Keycloak, manajemen user & role'],
    ])
    doc.add_paragraph('')
    doc.add_paragraph('Komunikasi antar layanan: NATS JetStream (event bisnis durable), HTTP internal (validasi lintas domain), dan Redis (cache replay SSE).')
    doc.add_page_break()

    # ===== 4. TEKNOLOGI =====
    add_colored_heading(doc, '4. Teknologi yang Digunakan', 1)
    add_table(doc, ['Komponen', 'Teknologi', 'Versi', 'Fungsi'], [
        ['Web Publik', 'Next.js + React + TypeScript', '16.3 / 19.2', 'Server-Side Rendering, PWA, SEO'],
        ['Web Admin', 'React + TypeScript + Vite', '19.x / 8.1', 'Single Page App dashboard operator'],
        ['Desain UI', 'Techwind 3.3.0 + Tailwind CSS', 'v4.x', 'Tema tunggal UI/UX + Dark Mode'],
        ['Backend', 'Golang', '1.26.6', 'Microservices HTTP & event handler'],
        ['Database', 'PostgreSQL + PostGIS', '15', 'Database per service, 7 database terisolasi'],
        ['Cache', 'Redis', '7 Alpine', 'Cache SSE replay, session, rate limit'],
        ['Message Broker', 'NATS JetStream', '2.10', 'Event bisnis durable (audit, livescore, medali)'],
        ['Identity', 'Keycloak', '24.0.0', 'SSO OpenID Connect, PKCE S256, realm roles'],
        ['Reverse Proxy', 'Nginx', '1.27 Alpine', 'SSL termination, CSP, rate limiting, routing'],
        ['Kontainerisasi', 'Docker + Docker Compose', 'Latest', '26 kontainer, named volumes, bridge network'],
        ['Observabilitas', 'Prometheus + Grafana', '2.45 / 10.0', 'Metrics scraping dan dashboard monitoring'],
        ['Peta Interaktif', 'Leaflet + React-Leaflet', 'Latest', 'Peta venue dengan tile OpenStreetMap'],
        ['Animasi', 'Framer Motion', 'Latest', 'Transisi halaman dan efek visual'],
        ['PWA', 'Serwist (Service Worker)', 'Latest', 'Instalasi, precaching, navigation preload'],
    ])
    doc.add_page_break()

    # ===== 5. WEB PUBLIK =====
    add_colored_heading(doc, '5. Web Publik (Public Portal) — Fitur Detail', 1)
    doc.add_paragraph('Dapat diakses tanpa login di https://porprov.depok.go.id. Berikut halaman dan fitur lengkapnya:')

    add_colored_heading(doc, '5.1 Beranda (Landing Page)', 2)
    nb(doc, 'Hero Banner Dinamis: Judul besar dengan teks sorotan berwarna gradient, latar belakang gambar parallax, badge tanggal "7 - 23 November 2026", dan hitung mundur (Hari, Jam, Menit, Detik) menuju Opening Ceremony.')
    nb(doc, 'Pengenalan PORPROV: Logo resmi, slogan "Bergerak Bersama Menuju Depok Maju!", 3 kartu pilar (Penjaringan Bakat, Pembinaan Atlet, Menuju Kompetisi Nasional).')
    nb(doc, 'Profil Maskot: Toca (Merah - Semangat Juang) dan Toci (Putih - Sportivitas).')
    nb(doc, '4 Kartu Pintasan: LiveScore, Jadwal, Medali, Cabor.')
    nb(doc, 'Venue Showcase: Grid 4 venue teratas dengan auto-refresh setiap 30 detik.')
    nb(doc, 'City Guide Bento Grid: Landmark Kota Depok (Tugu Selamat Datang, Alun-Alun, Kuliner, Akomodasi) + statistik (40+ Destinasi, 120+ Kuliner, 24/7 Transportasi).')

    add_colored_heading(doc, '5.2 Cabang Olahraga (/cabor)', 2)
    nb(doc, 'Halaman Indeks: Grid kartu seluruh cabor dengan ikon, nama, kategori, dan deskripsi singkat.')
    nb(doc, 'Halaman Detail (/cabor/[id]): Banner cabor, daftar nomor pertandingan, jadwal terkait, sidebar info teknis (Technical Delegate, total medali, status), dan venue terkait.')
    nb(doc, 'SEO: Metadata dinamis (title, description, OpenGraph) dibuat otomatis berdasarkan data cabor.')

    add_colored_heading(doc, '5.3 Venue Pertandingan (/venue)', 2)
    nb(doc, 'Halaman Indeks: Layout 2 kolom — Panel kiri: pencarian + daftar kartu venue. Panel kanan: Peta interaktif Leaflet dengan marker seluruh venue.')
    nb(doc, 'SSE Listener: Peta otomatis memperbarui data venue jika ada perubahan dari Admin (tanpa refresh).')
    nb(doc, 'Halaman Detail (/venue/[id]): Hero image, badge kesiapan arena, statistik (kapasitas, cabor, jadwal aktif), jadwal pertandingan di venue, tombol rute Google Maps, dan 6 rekomendasi destinasi terdekat (dihitung rumus Haversine dari koordinat).')

    add_colored_heading(doc, '5.4 Jadwal Pertandingan (/jadwal)', 2)
    nb(doc, 'Pencarian & Filter Komprehensif: Search bar kata kunci, 4 dropdown filter (Tanggal, Cabang Olahraga, Venue, Status).')
    nb(doc, 'Pengelompokan: Pertandingan dikelompokkan berdasarkan Cabang Olahraga.')
    nb(doc, 'Indikator: Waktu update terakhir dan tombol refresh manual.')
    nb(doc, 'Kartu Pertandingan: Menampilkan cabor, nomor tanding, peserta A vs B, venue, babak, status, dan waktu.')

    add_colored_heading(doc, '5.5 LiveScore Center (/livescore)', 2)
    nb(doc, 'Koneksi Ganda: Snapshot awal via REST API + pembaruan live via Server-Sent Events (SSE).')
    nb(doc, 'Validasi Integritas: Skor hanya ditampilkan jika jadwal memiliki tepat 2 peserta (A & B) yang lengkap.')
    nb(doc, 'Visual Feedback: Kartu pertandingan yang menerima update berkedip dan membesar selama 1 detik.')
    nb(doc, 'Status Bar: Menampilkan status koneksi realtime (Terhubung / Mencoba ulang), jumlah update, dan tipe transport (SSE + JetStream).')
    nb(doc, 'Edge Caching: Nginx secara eksplisit menetapkan Cache-Control: no-store untuk halaman ini.')

    add_colored_heading(doc, '5.6 Klasemen Medali (/medali)', 2)
    nb(doc, '3 Kartu Podium: Menyorot peringkat 1 (aksen emas), 2, dan 3.')
    nb(doc, 'Tabel Klasemen: Kolom Peringkat, Logo + Nama Kontingen, Emas, Perak, Perunggu, Total.')
    nb(doc, 'Pengurutan: Emas DESC > Perak DESC > Perunggu DESC > Nama ASC (bahasa Indonesia).')
    nb(doc, 'Pembaruan: SSE listener event MEDAL_STANDING_UPDATED + polling 30 detik sebagai fallback.')
    nb(doc, 'Hanya data medali berstatus OFFICIAL yang ditampilkan.')

    add_colored_heading(doc, '5.7 City Guide (/city-guide)', 2)
    nb(doc, '9 Kategori Filter: Semua, Coffee Shop, Wisata Kuliner, Tempat Menginap, Wisata Buatan, Wisata Situ, Pusat Perbelanjaan, Rumah Sakit, Lainnya.')
    nb(doc, 'Grid Kartu: Badge kategori, ikon dinamis, alamat, tombol rute Google Maps.')
    nb(doc, 'Pagination: 12 item per halaman, URL-synced (?category=&page=).')

    add_colored_heading(doc, '5.8 Fitur SEO & PWA', 2)
    nb(doc, 'Template Title: "%s | PORPROV XV 2026". OpenGraph & Twitter Card lengkap.')
    nb(doc, 'Sitemap XML: Prioritas Homepage 1.0, LiveScore 0.8 (frekuensi hourly), halaman lain daily.')
    nb(doc, 'Robots.txt: Mengizinkan /, memblokir /admin/, /api/, /realms/.')
    nb(doc, 'PWA Manifest: Display standalone, icon 192x192 & 512x512 maskable, Service Worker dengan precaching dan navigation preload.')
    doc.add_page_break()

    # ===== 6. WEB ADMIN =====
    add_colored_heading(doc, '6. Web Admin (Back-Office) — Fitur Detail', 1)
    doc.add_paragraph('Diakses di https://porprov.depok.go.id/admin/ dengan autentikasi Keycloak SSO.')

    add_colored_heading(doc, '6.1 Dashboard Overview', 2)
    doc.add_paragraph('Menampilkan 8 kartu statistik operasional:')
    add_table(doc, ['No', 'Kartu', 'Sumber Data', 'Ikon'], [
        ['1', 'Total Atlet Terdaftar', 'Placeholder (0)', 'Users'],
        ['2', 'Medali Didistribusikan', 'Placeholder (0)', 'Trophy'],
        ['3', 'Skor Masuk (Hari Ini)', 'Placeholder (0)', 'Activity'],
        ['4', 'Insiden Sistem', 'Placeholder (0)', 'AlertTriangle'],
        ['5', 'Total Cabang Olahraga', 'Real-time API /master-data/cabors', 'Medal'],
        ['6', 'Total Venue', 'Real-time API /venues', 'MapPin'],
        ['7', 'Total City Guide', 'Real-time API /master-data/city-guides', 'Map'],
        ['8', 'Total Kontingen', 'Real-time API /master-data/kontingens', 'Flag'],
    ])
    doc.add_paragraph('')
    nb(doc, 'Tabel Pertandingan Berlangsung: Waktu, Cabor, Pertandingan, Status.')
    nb(doc, '10 Log Sistem Terkini: Dari Audit Log, dengan indikator warna (Hijau: CREATE, Kuning: UPDATE, Merah: DELETE).')

    add_colored_heading(doc, '6.2 Recycle Bin (Pemulihan Data Terhapus)', 2)
    doc.add_paragraph('Menggabungkan data terhapus dari 3 endpoint: Master Data, Venue, dan Schedule.')
    nb(doc, 'Entitas yang didukung: Cabor, Nomor Tanding, Kontingen, City Guide, Media, Hero, Venue, dan Jadwal.')
    nb(doc, 'Setiap penghapusan meminta alasan (default: "Diarsipkan melalui Admin Web").')
    nb(doc, 'Tabel: Nama Data, Jenis Entitas, Waktu Diarsipkan, Actor & Alasan.')
    nb(doc, 'Tombol "Pulihkan": Mengembalikan data beserta seluruh relasinya.')

    add_colored_heading(doc, '6.3 Media Library', 2)
    nb(doc, 'Upload: Maksimal 10 MB, format JPG/PNG/WebP, nama file acak 16 byte.')
    nb(doc, 'Grid View: Thumbnail persegi, nama file, ukuran dalam KB.')
    nb(doc, 'Aksi: Salin URL ke clipboard, Arsipkan (soft delete dengan alasan).')
    nb(doc, 'Media Selector Modal: Digunakan di seluruh modul (Cabor, Venue, Hero, City Guide, Kontingen) sebagai picker gambar portabel.')

    add_colored_heading(doc, '6.4 Audit Log', 2)
    nb(doc, 'Pencatatan immutable (tidak dapat diedit/dihapus) setiap perubahan data sistem.')
    nb(doc, 'Kolom: Waktu, Versi Event, Aktor, Aksi, Service, Entitas, ID, Correlation (Event ID, Request ID), Integritas (SHA-256 Hash), Detail Payload (expandable JSON).')
    nb(doc, 'Filter: Pencarian teks, filter aksi, filter service, limit & offset.')
    nb(doc, 'Export CSV: Mengunduh seluruh log dalam format spreadsheet.')

    add_colored_heading(doc, '6.5 Manajemen Akun (User Management)', 2)
    nb(doc, 'Hanya dapat diakses oleh role super_admin.')
    nb(doc, 'Tambah User: Username, Email, Nama Lengkap, Role (dari Keycloak), Password.')
    nb(doc, 'Edit User: Sinkronisasi 2 arah dengan Keycloak (profil, password, role).')
    nb(doc, 'Hapus User: Menonaktifkan di Keycloak + soft-delete di database lokal.')
    nb(doc, 'Tabel: Username, Nama Lengkap, Email, Role, Tanggal Terdaftar.')

    add_colored_heading(doc, '6.6 Profil Akun', 2)
    nb(doc, 'Informasi: Username, Email, daftar Role (badge chip).')
    nb(doc, 'Tombol "Manajemen Akun": Membuka Keycloak Account Console untuk ganti password.')
    doc.add_page_break()

    # ===== 7. BACKEND =====
    add_colored_heading(doc, '7. Backend Microservices — Detail Teknis', 1)

    add_colored_heading(doc, '7.1 API Gateway — Gerbang Keamanan Tunggal', 2)
    nb(doc, 'Validasi JWT: Format Bearer, signature JWKS Keycloak, issuer check, expiration + 30 detik leeway, verifikasi azp/aud.')
    nb(doc, 'Propagasi Aktor: Menghapus header X-Actor-ID dari request luar, menyuntikkan kembali dari klaim sub JWT yang valid.')
    nb(doc, 'Error Masking: Menyaring response >= 500 dari downstream menjadi pesan stabil tanpa detail implementasi.')
    nb(doc, 'CORS Tunggal: Satu-satunya pengelola CORS, membersihkan header ganda dari downstream.')
    nb(doc, 'Batas Payload: 12 MB per request.')

    add_colored_heading(doc, '7.2 Pola Transactional Outbox (LiveScore & Medali)', 2)
    doc.add_paragraph('Untuk menjamin konsistensi data dan event, LiveScore Service dan Medal Standing Service menggunakan pola Transactional Outbox:')
    doc.add_paragraph('1. Perubahan data dan event disimpan dalam satu transaksi database PostgreSQL.', style='List Number')
    doc.add_paragraph('2. Worker latar belakang membaca event pending dari tabel outbox_events.', style='List Number')
    doc.add_paragraph('3. Event dipublikasikan ke NATS JetStream. Jika gagal, exponential backoff diterapkan.', style='List Number')
    doc.add_paragraph('4. Setelah berhasil, kolom published_at ditandai. Event tidak pernah hilang.', style='List Number')

    add_colored_heading(doc, '7.3 Immutability Database (WORM)', 2)
    doc.add_paragraph('Tiga tabel dilindungi oleh trigger PostgreSQL yang melempar exception pada UPDATE/DELETE:')
    nb(doc, 'livescore_revisions: Riwayat skor append-only.')
    nb(doc, 'audit_logs: Log mutasi permanen dengan SHA-256 hash integritas.')
    nb(doc, 'medal_submission_history: Riwayat transisi status medali.')

    add_colored_heading(doc, '7.4 Validasi Lintas Service', 2)
    nb(doc, 'Hapus Cabor: Memeriksa apakah ada Nomor Tanding aktif. Ditolak (409) jika ada.')
    nb(doc, 'Hapus Nomor Tanding: Memeriksa ke Schedule Service apakah ada jadwal aktif. Ditolak (409) jika ada.')
    nb(doc, 'Hapus Venue: Memeriksa ke Schedule Service apakah ada jadwal aktif. Ditolak (409) jika ada.')
    nb(doc, 'Buat Jadwal: Memvalidasi keberadaan Nomor Tanding, Kontingen, dan Venue di service masing-masing.')
    nb(doc, 'Submit Medali: Memvalidasi keberadaan Kontingen di Master Data Service.')
    nb(doc, 'Submit Skor: Memvalidasi keberadaan pertandingan dan kelengkapan 2 peserta di Schedule Service.')
    doc.add_page_break()

    # ===== 8. INFRASTRUKTUR =====
    add_colored_heading(doc, '8. Infrastruktur & Keamanan', 1)

    add_colored_heading(doc, '8.1 Docker Compose (26 Kontainer)', 2)
    doc.add_paragraph('Seluruh sistem dijalankan dengan satu perintah docker compose. Komposisi:')
    nb(doc, '1 PostgreSQL (7 database terisolasi), 1 Redis, 1 NATS JetStream, 1 Keycloak, 1 Nginx.')
    nb(doc, '7 migrasi database ephemeral (selesai lalu berhenti).')
    nb(doc, '9 layanan mikro Golang (user-service sampai realtime-gateway).')
    nb(doc, '1 API Gateway, 1 Public Web (Next.js), 1 Admin Web (React/Nginx).')
    nb(doc, '1 Keycloak Bootstrap (membuat realm, client, role, dan user awal).')
    nb(doc, '1 Prometheus + 1 Grafana (observabilitas).')

    add_colored_heading(doc, '8.2 Keamanan VPS Production', 2)
    nb(doc, 'SSL/TLS: Wildcard certificate depok.go.id, TLS 1.2/1.3, HTTP/2.')
    nb(doc, 'Docker Hardening: read_only filesystem, cap_drop ALL, no-new-privileges, tmpfs terbatas.')
    nb(doc, 'Rate Limiting: API 30 req/detik (burst 60), Auth 5 req/detik (burst 20).')
    nb(doc, 'Security Headers: HSTS (2 tahun), CSP dinamis per-path, X-Frame-Options, X-Content-Type-Options, COOP/COEP/CORP, Permissions-Policy.')
    nb(doc, 'CSP Engine: Nginx mengevaluasi CSP upstream secara bertingkat. Jika tidak lengkap (tanpa base-uri/frame-ancestors/form-action), fallback deterministik diterapkan.')
    nb(doc, 'Non-root Containers: Semua layanan Go berjalan sebagai user "porprov", Next.js sebagai "nextjs" (UID 1001).')

    add_colored_heading(doc, '8.3 Keycloak SSO', 2)
    nb(doc, 'Realm: "porprov" dengan 5 realm role dan 3 OIDC client.')
    nb(doc, 'Admin Web Client: Public client, Authorization Code + PKCE S256, tanpa Direct Access Grants.')
    nb(doc, 'Mobile Client: Public client, Standard Flow + Direct Access Grants untuk Expo.')
    nb(doc, 'Backend Service Client: Confidential, Service Account dengan manage-users & view-users.')
    nb(doc, 'User Awal: admin_depok (super_admin) dan koresponden_1 (koresponden).')
    doc.add_page_break()

    # ===== 9. RBAC =====
    add_colored_heading(doc, '9. Pemetaan Hak Akses (RBAC)', 1)
    add_table(doc, ['Menu / Fitur', 'super_admin', 'auditor', 'admin_venue', 'koresponden', 'verifikator'], [
        ['Dashboard', '✅', '✅', '✅', '✅', '✅'],
        ['Master Data', '✅', '✅', '✅', '✅', '✅'],
        ['Hero Utama', '✅', '✅', '✅', '✅', '✅'],
        ['LiveScore Center', '✅', '❌', '❌', '✅', '❌'],
        ['Perolehan Medali', '✅', '❌', '❌', '✅', '❌'],
        ['City Guide', '✅', '✅', '✅', '✅', '✅'],
        ['Media Library', '✅', '✅', '✅', '✅', '✅'],
        ['Verifikasi Medali', '✅', '❌', '❌', '❌', '✅'],
        ['Audit Log', '✅', '✅', '❌', '❌', '❌'],
        ['Manajemen Akun', '✅', '❌', '❌', '❌', '❌'],
        ['Profil Akun', '✅', '✅', '✅', '✅', '✅'],
        ['Medali: Submit', '✅', '❌', '❌', '✅', '❌'],
        ['Medali: Verify/Reject', '✅', '❌', '❌', '❌', '✅'],
        ['Medali: Publish', '✅', '❌', '❌', '❌', '❌'],
        ['SSE Stream Admin', '✅', '✅', '❌', '✅', '✅'],
    ])
    doc.add_page_break()

    # ===== 10. REKOMENDASI =====
    add_colored_heading(doc, '10. Rekomendasi Fitur Tambahan', 1)
    add_table(doc, ['No', 'Fitur', 'Manfaat', 'Prioritas'], [
        ['1', 'Aplikasi Mobile (React Native)', 'Pengalaman native di Android/iOS, fitur kamera untuk scan QR atlet', 'Tinggi'],
        ['2', 'Push Notification & WhatsApp', 'Notifikasi otomatis ke panitia saat ada pengajuan medali atau perubahan jadwal', 'Tinggi'],
        ['3', 'Peta Venue Interaktif Terpadu', 'Satu peta digital menampilkan seluruh venue beserta kepadatan dan cuaca', 'Sedang'],
        ['4', 'Kemitraan/Sponsor Corner', 'Spot iklan dinamis di Web Publik, dikelola dari Admin', 'Sedang'],
        ['5', 'Dashboard Analitik Lanjutan', 'Grafik tren skor, statistik atlet, heat map pertandingan', 'Rendah'],
        ['6', 'Sistem Tiket Digital', 'Penjualan dan validasi tiket penonton secara elektronik', 'Rendah'],
    ])
    doc.add_page_break()

    # ===== 11. KEBUTUHAN =====
    add_colored_heading(doc, '11. Kebutuhan Sistem', 1)
    add_table(doc, ['Komponen', 'Spesifikasi Minimal', 'Spesifikasi Rekomendasi'], [
        ['Sistem Operasi', 'Ubuntu 22.04 LTS', 'Ubuntu 24.04 LTS'],
        ['RAM', '8 GB', '16 GB'],
        ['CPU', '4 vCPU', '8 vCPU'],
        ['Storage', '50 GB SSD', '100 GB NVMe SSD'],
        ['Docker', 'Docker Engine 24+', 'Docker Engine 27+'],
        ['Domain', 'Subdomain Diskominfo', 'porprov.depok.go.id (sudah aktif)'],
        ['SSL', 'Let\'s Encrypt', 'Wildcard SSL depok.go.id (sudah terpasang)'],
        ['Bandwidth', '10 Mbps', '100 Mbps (untuk hari pertandingan)'],
    ])

    doc.add_paragraph('')
    doc.add_paragraph('')
    p_end = doc.add_paragraph(
        'Dokumen ini disusun berdasarkan pembedahan langsung seluruh kode sumber '
        'repository porprov-depok, konfigurasi infrastruktur Docker, skema database, '
        'dan arsitektur layanan mikro yang telah berjalan di VPS Publik '
        'https://porprov.depok.go.id. Seluruh informasi akurat per Agustus 2026.',
        style='Quote')
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output = Path(__file__).resolve().parent / "Laporan_Analisis_Komprehensif_Detail_Porprov_XV.docx"
    doc.save(output)
    print(f"Document saved to {output}")

if __name__ == "__main__":
    main()
