import os
from docx import Document
from docx.shared import Pt, Inches

def create_wireframe_docx():
    doc = Document()
    doc.add_heading('ASCII Wireframe Portal PORPROV XV Jawa Barat 2026 Kota Depok', 0)
    
    doc.add_heading('1. Prinsip ASCII Wireframe', level=1)
    doc.add_paragraph('Dokumen ini adalah rancangan tampilan dan alur interaksi yang diselaraskan dengan arsitektur Techwind 3.3.0 untuk web publik dan React Dashboard untuk admin. Semua komponen diwajibkan mobile-first dan mengikuti standar WCAG 2.2 AA.')
    
    doc.add_heading('2. Struktur Navigasi Publik (Next.js PWA)', level=1)
    doc.add_paragraph('- Beranda: Hero parallax 50%, Ticker LiveScore, Standings Medali Preview\n- Cabor: Detail pertandingan dan hasil\n- Jadwal: Kalender dengan filter (Date, Cabor, Venue, Status)\n- Venue: Integrasi Maps, Rute, City Guide sekitar\n- LiveScore: Halaman realtime (Event-Driven dari NATS/Gateway)\n- Medali: Klasemen resmi dengan fallback polling\n- Depok Guide: Rekomendasi tempat hasil kurasi')
    
    doc.add_heading('3. Struktur Backend Admin (React Vite)', level=1)
    doc.add_paragraph('- Master Data: Cabor, Nomor, Kontingen, City Guide (Soft Delete)\n- Media Library: Integrasi volume persisten tunggal\n- LiveScore Center: Private SSE, verifikasi pertandingan\n- Medali: Alur PENDING -> VERIFIED -> OFFICIAL\n- Audit Log: Akses baca untuk Super Admin dan Auditor')
    
    doc.add_heading('4. Mockup (Layout)', level=1)
    p = doc.add_paragraph()
    p.add_run('+---------------------------------------------------------+\n').font.name = 'Courier New'
    p.add_run('| [LOGO] Beranda | Jadwal | LiveScore | Venue | Medali  |\n').font.name = 'Courier New'
    p.add_run('+---------------------------------------------------------+\n').font.name = 'Courier New'
    p.add_run('|                                                         |\n').font.name = 'Courier New'
    p.add_run('|         HERO SECTION (100vh)                            |\n').font.name = 'Courier New'
    p.add_run('|         Bergerak Bersama Menuju Depok Maju              |\n').font.name = 'Courier New'
    p.add_run('|                                                         |\n').font.name = 'Courier New'
    p.add_run('+---------------------------------------------------------+').font.name = 'Courier New'

    doc.save('design/Final_ASCII_Wireframe.docx')

def create_brd_docx():
    doc = Document()
    doc.add_heading('BRD PRD SRS SDD Terpadu - Portal PORPROV XV', 0)
    
    doc.add_heading('1. Business Requirements Document (BRD)', level=1)
    doc.add_paragraph('Sistem ini digunakan sebagai pusat informasi publik yang realtime serta portal administrasi panitia PORPROV XV. Objektif utamanya adalah publikasi jadwal akurat, LiveScore yang low-latency, serta informasi lokasi (venue) dan pariwisata terintegrasi.')
    
    doc.add_heading('2. Product Requirements Document (PRD)', level=1)
    doc.add_paragraph('Fitur utama meliputi:\n1. Public Web dengan performa SEO (LCP < 2.5s)\n2. Admin Dashboard dengan Role-Based Access Control dari Keycloak\n3. Mobile Koresponden dengan kapabilitas Offline Queue untuk input skor di lapangan\n4. Realtime Gateway via NATS JetStream\n5. Soft Delete data master dan Recycle Bin')
    
    doc.add_heading('3. Software Requirements Specification (SRS)', level=1)
    doc.add_paragraph('Kebutuhan sistem:\n- Frontend: React / Next.js / Tailwind CSS v4\n- Backend: Golang microservices\n- Database: PostgreSQL (per service)\n- Broker: NATS JetStream\n- Auth: Keycloak (OIDC/OAuth2)')
    
    doc.add_heading('4. Software Design Document (SDD)', level=1)
    doc.add_paragraph('Desain perangkat lunak mematuhi standar microservices dengan Database-per-service. LiveScore Service menyimpan histori append-only (livescore_revisions). Komunikasi antar service mengandalkan event NATS (transactional outbox) agar konsistensi data medali dan skor terjamin tanpa distributed transaction locking.')

    doc.save('design/Final_BRD_PRD_SRS_SDD.docx')

def create_architecture_docx():
    doc = Document()
    doc.add_heading('Dokumen Perencanaan Arsitektur Enterprise Web & Mobile', 0)
    
    doc.add_heading('1. Topologi dan Infrastruktur (Final)', level=1)
    doc.add_paragraph('Sistem dijalankan melalui Docker Compose canonical yang memisahkan database tiap core service (master_data_db, schedule_db, livescore_db, porprov_db, venue_db). Traffic luar masuk secara eksklusif dari API Gateway pada port 8000.')
    
    doc.add_heading('2. Standar Port (Namespace)', level=1)
    doc.add_paragraph('Alokasi port lokal (debugging go run) menggunakan awalan 28xxx:\n- 28000: API Gateway\n- 28001: User Service\n- 28081: Master Data Service\n- 28082: Schedule Service\n- 28083: LiveScore Service\n- 28084: Audit Service\n- 28085: Realtime Gateway\n- 28086: Medal Standing Service\n- 28087: Venue Service')
    
    doc.add_heading('3. Keamanan', level=1)
    doc.add_paragraph('Penggunaan API Gateway (memvalidasi token Keycloak). CORS headers eksklusif hanya diberikan oleh API Gateway. Setiap perubahan data yang berkaitan dengan jadwal dan pertandingan di-record ke dalam sistem Immutable Audit Logging.')
    
    doc.add_heading('4. Transactional Outbox Pattern', level=1)
    doc.add_paragraph('Modul LiveScore dan Medal menggunakan pendekatan Outbox Pattern, di mana tabel events ditulis bersamaan dengan update tabel state pada satu transaksi PostgreSQL, kemudian sebuah NATS publisher membaca outbox ini secara asinkron dan mengirimkannya (At-Least-Once delivery).')
    
    doc.save('design/Final_Perencanaan_Arsitektur.docx')

if __name__ == '__main__':
    print("Generating DOCX files...")
    create_wireframe_docx()
    create_brd_docx()
    create_architecture_docx()
    print("Done generating 3 DOCX files.")
