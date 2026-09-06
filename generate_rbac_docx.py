from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

def main():
    doc = Document()

    # Title
    title = doc.add_heading('Pemetaan Hak Akses (Role Mapping) - Web Admin PORPROV', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph(
        'Dokumen ini berisi daftar hak akses berdasarkan Role (Peran) '
        'untuk sistem Web Admin Portal PORPROV XV Jawa Barat 2026. '
        'Sistem ini mengimplementasikan Role-Based Access Control (RBAC) melalui Keycloak '
        'dan divalidasi terpusat pada API Gateway.'
    )

    roles = [
        {
            "role": "super_admin",
            "desc": "Administrator tertinggi yang memiliki kontrol penuh terhadap seluruh sistem.",
            "access": [
                "Full akses ke semua Modul.",
                "Modul Manajemen Pengguna (Tambah, Edit, Hapus User & Role).",
                "Modul Master Data (Cabor, Venue, Jadwal, Peserta).",
                "Modul Medali (Tahap final: PUBLISH medali ke klasemen publik).",
                "Modul LiveScore (Manajemen seluruh skor pertandingan).",
                "Modul Audit Log (Melihat jejak mutasi data sistem)."
            ]
        },
        {
            "role": "auditor",
            "desc": "Pengguna dengan peran pemantauan dan kepatuhan sistem.",
            "access": [
                "Modul Audit Log (Membaca seluruh jejak rekam mutasi data / Immutable Log).",
                "Membaca Stream Events (SSE) dari sistem internal.",
                "Akses Read-Only ke Master Data (sesuai kebutuhan audit)."
            ]
        },
        {
            "role": "admin_venue",
            "desc": "Administrator yang bertugas mengelola titik-titik venue dan fasilitas pertandingan.",
            "access": [
                "Modul Master Data -> Manajemen Venue.",
                "Modul Master Data -> City Guide (Destinasi/Fasilitas).",
                "Modul Master Data -> Cabang Olahraga."
            ]
        },
        {
            "role": "koresponden",
            "desc": "Petugas lapangan/reporter yang bertugas melaporkan hasil pertandingan secara real-time.",
            "access": [
                "Modul LiveScore (Memperbarui skor pertandingan secara Real-Time).",
                "Modul Medali (Tahap SUBMIT: Mengajukan data pemenang/medali).",
                "Melihat Jadwal Pertandingan dan Peserta."
            ]
        },
        {
            "role": "verifikator",
            "desc": "Petugas panel tingkat menengah yang memvalidasi keabsahan data lapangan.",
            "access": [
                "Modul Medali (Tahap VERIFY/REJECT: Memeriksa dan menyetujui pengajuan medali dari Koresponden).",
                "Membaca Stream Events internal (SSE) terkait pembaruan pertandingan.",
                "Melihat Jadwal Pertandingan dan Peserta."
            ]
        }
    ]

    doc.add_heading('Matriks Peran dan Hak Akses', level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Table header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role (Peran)'
    hdr_cells[1].text = 'Deskripsi'
    hdr_cells[2].text = 'Modul / Akses yang Diizinkan'

    # Bold header
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Populate table
    for role in roles:
        row_cells = table.add_row().cells

        # Role name
        p0 = row_cells[0].paragraphs[0]
        r0 = p0.add_run(role['role'])
        r0.font.bold = True
        r0.font.name = 'Courier New'

        # Description
        row_cells[1].text = role['desc']

        # Access List
        p2 = row_cells[2].paragraphs[0]
        for access in role['access']:
            p2.add_run(f"• {access}\n")

    # Conclusion
    doc.add_paragraph("\n")
    doc.add_heading('Catatan Tambahan', level=2)
    doc.add_paragraph(
        '1. Semua Role membutuhkan autentikasi (Login) melalui Keycloak.\n'
        '2. Pemisahan tugas yang tegas terlihat pada Modul Medali di mana pengajuan (Submit) hanya dapat dilakukan oleh koresponden, verifikasi (Verify) oleh verifikator, dan pengesahan akhir (Publish) hanya oleh super_admin.\n'
        '3. Log mutasi untuk setiap aksi tercatat secara abadi (immutable) dan hanya bisa dievaluasi oleh super_admin dan auditor.'
    )

    # Save doc
    output_path = Path(__file__).resolve().parent / "Pemetaan_Akses_Role_Porprov.docx"
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    main()
