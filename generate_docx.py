import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = docx.Document()

# Title
title = doc.add_heading('Katalog API & Panduan Integrasi\nPortal PORPROV XV Jawa Barat 2026', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Strategi Integrasi
doc.add_heading('1. Strategi Integrasi Antar-Sistem (B2B)', level=1)
p = doc.add_paragraph('Untuk menyalurkan (mengintegrasikan) API Gateway PORPROV dengan aplikasi/sistem eksternal di masa depan, sistem menggunakan mekanisme ')
p.add_run('Client Credentials Grant').bold = True
p.add_run(' melalui Keycloak (OAuth2.0). Langkah-langkahnya adalah sebagai berikut:')

steps = [
    "Pendaftaran Klien: Sistem eksternal didaftarkan sebagai 'Client' baru di Keycloak (misal: sistem-koni-jabar) dengan mode Confidential.",
    "Kredensial: Sistem eksternal akan mendapatkan Client ID dan Client Secret yang sangat rahasia.",
    "Otorisasi: Administrator memberikan 'Roles' atau 'Scopes' spesifik kepada klien tersebut (contoh: hanya bisa membaca LiveScore, tidak bisa mengubah).",
    "Permintaan Token: Sistem eksternal meminta JWT Access Token dengan menukarkan Client ID dan Secret ke Endpoint Token Keycloak.",
    "Akses API Gateway: Sistem eksternal memanggil API Gateway dengan menyertakan token JWT tersebut pada header 'Authorization: Bearer <token>'."
]
for step in steps:
    doc.add_paragraph(step, style='List Bullet')

# 2. Endpoint Autentikasi
doc.add_heading('2. Autentikasi (Keycloak)', level=1)
doc.add_heading('Dapatkan Access Token', level=2)
p = doc.add_paragraph()
p.add_run('Tautan API: ').bold = True
p.add_run('POST https://10.11.5.80/realms/porprov/protocol/openid-connect/token')
doc.add_paragraph('Deskripsi: Digunakan oleh sistem eksternal untuk mendapatkan JWT Access Token.')
p = doc.add_paragraph()
p.add_run('Cara Penggunaan:').bold = True
code = doc.add_paragraph('Header: Content-Type: application/x-www-form-urlencoded\nBody:\n  grant_type=client_credentials\n  client_id=<CLIENT_ID>\n  client_secret=<CLIENT_SECRET>')
# Macro text might not exist by default, just use normal
# code.style = 'Macro Text'

# 3. Katalog API Terbuka (Publik & Terotorisasi)
doc.add_heading('3. Katalog API Gateway Utama', level=1)
p = doc.add_paragraph()
p.add_run('Basis URL (Base URL): ').bold = True
p.add_run('https://10.11.5.80/api/v1')

apis = [
    {
        "name": "Daftar Cabang Olahraga (Master Data)",
        "endpoint": "GET /master-data/cabor",
        "desc": "Mengambil seluruh daftar cabang olahraga (Cabor) yang dipertandingkan beserta detailnya.",
        "usage": "Kirim HTTP GET ke endpoint. Tidak memerlukan token untuk baca publik. Jika dengan token, akan mengembalikan data lebih lengkap."
    },
    {
        "name": "Daftar Lokasi (Venue)",
        "endpoint": "GET /venues",
        "desc": "Mendapatkan daftar fasilitas olahraga dan peta koordinat GPS (City Guide).",
        "usage": "Kirim HTTP GET ke endpoint. Parameter opsional: ?limit=10&page=1"
    },
    {
        "name": "Jadwal Pertandingan Harian",
        "endpoint": "GET /schedules?date=YYYY-MM-DD",
        "desc": "Mendapatkan jadwal pertandingan terpadu untuk tanggal tertentu.",
        "usage": "Kirim HTTP GET dengan format tanggal ISO. Respons mencakup daftar pertandingan, venue, dan status selesai/belum."
    },
    {
        "name": "Realtime LiveScore (Server-Sent Events)",
        "endpoint": "GET /stream/events",
        "desc": "Terhubung dengan sistem NATS JetStream untuk mendapatkan pembaruan skor pertandingan secara real-time tanpa me-refresh halaman.",
        "usage": "Gunakan EventSource API pada client untuk mendengarkan event 'score_updated'. Akses publik akan menerima skor tersanitasi."
    },
    {
        "name": "Perolehan Medali (Medal Standing)",
        "endpoint": "GET /medals/standings",
        "desc": "Mendapatkan klasemen akhir perolehan medali (Emas, Perak, Perunggu) berdasarkan kontingen/kota.",
        "usage": "Kirim HTTP GET. Hanya mengembalikan data medali yang sudah berstatus VERIFIED (resmi)."
    }
]

for api in apis:
    doc.add_heading(api["name"], level=2)

    p = doc.add_paragraph()
    p.add_run('Tautan API: ').bold = True
    p.add_run(api["endpoint"])

    p = doc.add_paragraph()
    p.add_run('Deskripsi: ').bold = True
    p.add_run(api["desc"])

    p = doc.add_paragraph()
    p.add_run('Cara Penggunaan: ').bold = True
    p.add_run(api["usage"])

    doc.add_paragraph('')

# Save document
doc.save(Path(__file__).resolve().parent / 'Katalog_API_Porprov_XV.docx')
print("Katalog API DOCX berhasil dibuat!")
