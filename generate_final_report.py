import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# === HELPER FUNCTIONS ===
def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    s = tc.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tc.append(s)

def h1(doc, text): _h(doc, text, 1)
def h2(doc, text): _h(doc, text, 2)
def h3(doc, text): _h(doc, text, 3)
def _h(doc, text, level):
    hd = doc.add_heading(text, level)
    for r in hd.runs:
        r.font.color.rgb = RGBColor(0x13, 0x57, 0x9B)

def para(doc, text):
    doc.add_paragraph(text)

def bold_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p

def bullet(doc, title, text):
    p = doc.add_paragraph(style='List Bullet')
    if title:
        r = p.add_run(title)
        r.bold = True
    p.add_run(text)

def nbullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def num(doc, text):
    doc.add_paragraph(text, style='List Number')

def table(doc, headers, rows, header_color='13579B'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, header_color)
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
    for rd in rows:
        row = t.add_row()
        for i, val in enumerate(rd):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph('')
    return t

def stat_table(doc, items):
    """Create a highlight statistics row"""
    t = doc.add_table(rows=2, cols=len(items))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    colors = ['1B6B93', 'E8630A', '2E8B57', 'B8860B', '8B3A8B', '13579B']
    for i, (num_val, label) in enumerate(items):
        c_top = t.rows[0].cells[i]
        c_bot = t.rows[1].cells[i]
        c_top.text = str(num_val)
        c_bot.text = label
        shade(c_top, colors[i % len(colors)])
        shade(c_bot, 'F0F4F8')
        for p in c_top.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(18)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for p in c_bot.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(8)
                r.font.bold = True
    doc.add_paragraph('')

def divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('━' * 60)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(8)

def section_intro(doc, emoji, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ==================== MAIN ====================
def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10)
    style.font.name = 'Calibri'
    style.paragraph_format.space_after = Pt(4)

    # ============================================================
    #  COVER PAGE
    # ============================================================
    for _ in range(4): doc.add_paragraph('')

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run('DINAS KOMUNIKASI DAN INFORMATIKA')
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x13, 0x57, 0x9B)
    r.bold = True

    cp2 = doc.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cp2.add_run('KOTA DEPOK')
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x13, 0x57, 0x9B)
    r2.bold = True

    doc.add_paragraph('')
    divider(doc)
    doc.add_paragraph('')

    tt = doc.add_paragraph()
    tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tt.add_run('LAPORAN ANALISIS KOMPREHENSIF')
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x13, 0x57, 0x9B)

    tt2 = doc.add_paragraph()
    tt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tt2.add_run('Portal Digital PORPROV XV\nJawa Barat 2026')
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph('')
    divider(doc)
    doc.add_paragraph('')

    for line in ['https://porprov.depok.go.id', '', 'Versi Dokumen: 1.0', 'Agustus 2026']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ============================================================
    #  DAFTAR ISI
    # ============================================================
    h1(doc, 'Daftar Isi')
    toc = [
        'BAB I    — Selayang Pandang & Angka Kunci',
        'BAB II   — Visi Arsitektur: Mengapa Microservices?',
        'BAB III  — Teknologi di Balik Layar',
        'BAB IV   — Alur Kerja Lapangan (Proses Bisnis)',
        'BAB V    — Fitur Web Publik: Pengalaman Masyarakat',
        'BAB VI   — Fitur Web Admin: Senjata Panitia',
        'BAB VII  — Keamanan & Ketahanan Sistem',
        'BAB VIII — Pemetaan Hak Akses (RBAC)',
        'BAB IX   — Infrastruktur & Deployment',
        'BAB X    — Rekomendasi Pengembangan Selanjutnya',
        'BAB XI   — Kebutuhan Sistem & Spesifikasi',
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        r = p.runs[0]
        r.font.size = Pt(11)
    doc.add_page_break()

    # ============================================================
    #  BAB I — SELAYANG PANDANG
    # ============================================================
    h1(doc, 'BAB I — Selayang Pandang & Angka Kunci')
    section_intro(doc, '', 'Satu portal, satu visi: menyatukan informasi pertandingan PORPROV XV dari meja panitia hingga layar gawai masyarakat Jawa Barat.')
    doc.add_paragraph('')

    para(doc, 'Portal PORPROV XV Jawa Barat 2026 adalah sistem informasi pertandingan olahraga tingkat provinsi yang dibangun secara khusus untuk Kota Depok sebagai tuan rumah. Sistem ini menghubungkan dua dunia yang berbeda secara mulus: dunia operasional panitia (Web Admin) dan dunia informasi publik (Web Publik) — keduanya berjalan di atas satu ekosistem terintegrasi.')

    para(doc, 'Berikut adalah potret cepat kapasitas sistem yang saat ini telah beroperasi:')

    stat_table(doc, [
        ('26', 'Kontainer\nDocker'),
        ('9', 'Layanan\nMikro'),
        ('7', 'Database\nTerisolasi'),
        ('5', 'Peran\nPengguna'),
        ('60+', 'Endpoint\nAPI'),
        ('14', 'Halaman\nWeb'),
    ])

    h2(doc, 'Siapa Saja Pengguna Sistem Ini?')
    table(doc, ['Pengguna', 'Akses', 'Kegiatan Utama'], [
        ['Masyarakat Umum', 'Web Publik (tanpa login)', 'Melihat jadwal, skor langsung, klasemen medali, dan panduan wisata Kota Depok'],
        ['Koresponden (Petugas Lapangan)', 'Web Admin (login)', 'Memasukkan skor pertandingan secara real-time dan mengajukan perolehan medali'],
        ['Verifikator (Panel Juri)', 'Web Admin (login)', 'Memverifikasi atau menolak pengajuan medali dari koresponden'],
        ['Admin Venue / Organisasi', 'Web Admin (login)', 'Mengelola data cabor, venue, kontingen, city guide, dan jadwal pertandingan'],
        ['Auditor', 'Web Admin (login)', 'Memantau seluruh jejak perubahan data sistem (read-only)'],
        ['Super Admin', 'Web Admin (login)', 'Mengelola semua fitur termasuk manajemen akun dan pengesahan medali resmi'],
    ])
    doc.add_page_break()

    # ============================================================
    #  BAB II — ARSITEKTUR
    # ============================================================
    h1(doc, 'BAB II — Visi Arsitektur: Mengapa Microservices?')
    section_intro(doc, '', 'Bayangkan sebuah orkestra: setiap musisi memainkan instrumennya secara mandiri, namun bersama-sama menghasilkan simfoni yang utuh.')
    doc.add_paragraph('')

    para(doc, 'Sistem ini tidak dibangun sebagai satu blok besar (monolith) yang rentan. Sebaliknya, ia dipecah menjadi 9 layanan mikro yang masing-masing bertanggung jawab atas satu domain bisnis. Jika satu layanan mengalami beban berat, layanan lain tetap beroperasi normal — persis seperti orkestra yang tetap bermain meski satu pemain istirahat.')

    h2(doc, 'Peta 9 Layanan Mikro')
    table(doc, ['No', 'Layanan', 'Tanggung Jawab', 'Database'], [
        ['1', 'API Gateway', 'Gerbang tunggal: validasi identitas, otorisasi, routing, penyaringan error', '— (tanpa DB)'],
        ['2', 'Master Data', 'Cabor, Kontingen, Nomor Tanding, City Guide, Media Library, Hero', 'master_data_db'],
        ['3', 'Schedule', 'Jadwal pertandingan dan susunan Peserta A vs B', 'schedule_db'],
        ['4', 'Venue', 'Lokasi GOR/Stadion, koordinat peta, kapasitas, fasilitas', 'venue_db'],
        ['5', 'LiveScore', 'Skor real-time, revisi append-only, koreksi skor', 'livescore_db'],
        ['6', 'Medal Standing', 'Workflow medali 4 tahap dan klasemen resmi', 'porprov_db'],
        ['7', 'Audit', 'Log mutasi immutable (tidak bisa diubah/dihapus)', 'audit_db'],
        ['8', 'Realtime Gateway', 'Pemancar SSE ke browser publik & admin', 'Redis (cache)'],
        ['9', 'User Service', 'Sinkronisasi akun Keycloak & manajemen pengguna', 'user_service_db'],
    ])

    h2(doc, 'Bagaimana Mereka Berkomunikasi?')
    table(doc, ['Jalur Komunikasi', 'Teknologi', 'Contoh Penggunaan'], [
        ['Event Bisnis (Asinkron)', 'NATS JetStream', 'Skor diperbarui → event dikirim → Audit Log tercatat otomatis'],
        ['Validasi Lintas Domain (Sinkron)', 'HTTP Internal', 'Hapus Cabor → cek ke Schedule apakah ada jadwal aktif'],
        ['Cache & Replay', 'Redis', 'Koneksi SSE baru → langsung mendapat snapshot skor terkini'],
        ['Identitas Terpusat', 'Keycloak OIDC', 'Login satu kali → akses ke seluruh sistem (Single Sign-On)'],
    ])
    doc.add_page_break()

    # ============================================================
    #  BAB III — TEKNOLOGI
    # ============================================================
    h1(doc, 'BAB III — Teknologi di Balik Layar')
    section_intro(doc, '', 'Fondasi teknologi yang sama dengan yang digunakan oleh perusahaan teknologi kelas dunia.')
    doc.add_paragraph('')

    table(doc, ['Lapisan', 'Teknologi', 'Versi', 'Mengapa Dipilih?'], [
        ['Tampilan Publik', 'Next.js + React + TypeScript', '16.3 / 19.2', 'SSR untuk SEO, PWA untuk instalasi, responsif di semua perangkat'],
        ['Tampilan Admin', 'React + TypeScript + Vite', '19.x / 8.1', 'SPA cepat untuk operasional panitia sehari-hari'],
        ['Tema UI/UX', 'Techwind 3.3.0 + Tailwind CSS v4', '3.3 / 4.x', 'Tema tunggal yang konsisten, mendukung Dark Mode, WCAG 2.2 AA'],
        ['Backend', 'Golang (Go)', '1.26.6', 'Super cepat, irit memori, jago menangani ribuan request bersamaan'],
        ['Database Utama', 'PostgreSQL + PostGIS', '15', 'Andal, mendukung transaksi ACID dan data geospasial'],
        ['Cache Kilat', 'Redis', '7 Alpine', 'Penyimpan sementara untuk replay SSE dan rate limiting'],
        ['Pesan Antarservice', 'NATS JetStream', '2.10', 'Event bisnis durable yang tidak pernah hilang'],
        ['Identitas (SSO)', 'Keycloak', '24.0', 'OpenID Connect, PKCE S256, manajemen role terpusat'],
        ['Reverse Proxy', 'Nginx', '1.27', 'SSL termination, rate limiting, Content Security Policy'],
        ['Kontainerisasi', 'Docker Compose', 'Latest', '26 kontainer dijalankan dengan satu perintah'],
        ['Peta Interaktif', 'Leaflet + OpenStreetMap', 'Latest', 'Peta venue tanpa biaya lisensi'],
        ['Monitoring', 'Prometheus + Grafana', '2.45 / 10.0', 'Dashboard monitoring kesehatan sistem'],
    ])
    doc.add_page_break()

    # ============================================================
    #  BAB IV — PROSES BISNIS
    # ============================================================
    h1(doc, 'BAB IV — Alur Kerja Lapangan (Proses Bisnis)')
    section_intro(doc, '', 'Dari meja kantor hingga pinggir lapangan — begini perjalanan data dari awal hingga skor terpampang di layar warga.')
    doc.add_paragraph('')

    # --- TAHAP 1 ---
    h2(doc, 'TAHAP 1: Membangun Pondasi Data')
    para(doc, 'Langkah pertama sebelum pertandingan bisa dijadwalkan. Dilakukan oleh Admin Organisasi melalui Web Admin.')

    h3(doc, '1a. Mendaftarkan Cabang Olahraga')
    para(doc, 'Menu: Master Data → Cabang Olahraga → Tombol "Tambah"')
    table(doc, ['Isian Form', 'Jenis Input', 'Penjelasan'], [
        ['Nama Cabor', 'Teks (wajib)', 'Misal: "Sepak Bola", "Pencak Silat", "Renang"'],
        ['Logo Cabor', 'Pilih dari Media Library', 'Gambar ikon resmi cabor'],
        ['Kategori', 'Dropdown', 'Tanding / Seni-Terukur / E-Sports / Eksibisi'],
        ['Status', 'Dropdown', 'Aktif / Eksibisi / Non-Aktif'],
        ['Total Medali', 'Angka', 'Jumlah medali yang diperebutkan di cabor ini'],
        ['Technical Delegate', 'Teks', 'Nama penanggung jawab teknis'],
        ['Deskripsi', 'Teks Panjang', 'Keterangan tambahan'],
    ])

    h3(doc, '1b. Mendaftarkan Nomor Pertandingan')
    para(doc, 'Setiap cabor memiliki sub-kategori bertanding. Menu: Master Data → Nomor Pertandingan → "Tambah"')
    table(doc, ['Isian Form', 'Jenis Input', 'Penjelasan'], [
        ['Cabang Olahraga', 'Dropdown pencarian (wajib)', 'Memilih cabor induk yang sudah terdaftar'],
        ['Nama Nomor', 'Teks (wajib)', 'Misal: "Tunggal Putra", "Ganda Campuran"'],
        ['Kategori Gender', 'Dropdown', 'Putra / Putri / Campuran / Terbuka'],
        ['Tipe Pertandingan', 'Dropdown', 'Tanding / Seni / Terukur / Beregu'],
    ])

    h3(doc, '1c. Mendaftarkan Kontingen Peserta')
    para(doc, 'Setiap Kota/Kabupaten di Jawa Barat didaftarkan sebagai kontingen. Menu: Master Data → Kontingen → "Tambah"')
    table(doc, ['Isian Form', 'Jenis Input', 'Penjelasan'], [
        ['Nama Kontingen', 'Teks (wajib)', 'Misal: "Kota Depok", "Kabupaten Bogor"'],
        ['Tipe Daerah', 'Dropdown', 'Kota atau Kabupaten'],
        ['Logo Kontingen', 'Pilih dari Media Library', 'Logo resmi daerah (opsional)'],
    ])

    h3(doc, '1d. Mendaftarkan Venue Pertandingan')
    para(doc, 'Setiap GOR/Stadion tempat bertanding didaftarkan. Menu: Master Data → Venue → "Tambah"')
    table(doc, ['Isian Form', 'Jenis Input', 'Penjelasan'], [
        ['Nama Venue', 'Teks (wajib)', 'Misal: "GOR Kartika Kostrad"'],
        ['Status Kesiapan', 'Dropdown', 'Persiapan / Siap / Sedang Digunakan'],
        ['Kapasitas Penonton', 'Angka (wajib)', 'Jumlah kursi penonton'],
        ['Cabor Terkait', 'Multi-select tag', 'Cabor yang akan bertanding di venue ini'],
        ['Alamat Lengkap', 'Teks Panjang', 'Alamat fisik venue'],
        ['Latitude', 'Desimal (wajib)', 'Koordinat lintang, rentang -90 s.d. 90'],
        ['Longitude', 'Desimal (wajib)', 'Koordinat bujur, rentang -180 s.d. 180'],
        ['URL Google Maps', 'URL HTTPS', 'Link rute resmi Google Maps (opsional)'],
        ['Gambar Venue', 'Pilih dari Media Library', 'Foto venue'],
        ['Contact Person', 'Teks', 'Nama / No HP penanggung jawab'],
        ['Fasilitas', 'Teks', 'Misal: Toilet, Parkir, Ruang Medis'],
    ])

    h3(doc, '1e. Mendaftarkan City Guide (Panduan Wisata)')
    para(doc, 'Destinasi wisata, kuliner, dan fasilitas Kota Depok. Menu: City Guide → "Tambah"')
    table(doc, ['Isian Form', 'Jenis Input', 'Penjelasan'], [
        ['Judul', 'Teks (wajib, maks 255)', 'Nama destinasi'],
        ['Kategori', 'Dropdown', '8 pilihan: Coffee Shop, Wisata Kuliner, Tempat Menginap, Wisata Buatan, Wisata Situ, Pusat Perbelanjaan, Rumah Sakit, Lainnya'],
        ['Deskripsi', 'Teks Panjang', 'Keterangan tentang destinasi'],
        ['Alamat', 'Teks Panjang', 'Alamat fisik'],
        ['URL Google Maps', 'URL HTTPS', 'Link rute (maks 2048 karakter)'],
        ['Latitude & Longitude', 'Desimal (wajib)', 'Koordinat pasangan untuk integrasi peta'],
        ['Tombol GPS', 'Tombol Aksi', 'Mengambil koordinat otomatis dari perangkat'],
        ['Gambar', 'Pilih dari Media Library', 'Foto destinasi'],
    ])
    doc.add_page_break()

    # --- TAHAP 2 ---
    h2(doc, 'TAHAP 2: Menyusun Jadwal Pertandingan')
    para(doc, 'Setelah pondasi lengkap, Admin merakit jadwal. Menu: Master Data → Jadwal Pertandingan → "Tambah"')
    table(doc, ['Isian Form', 'Jenis Input', 'Penjelasan'], [
        ['Cabor / Nomor Tanding', 'Dropdown pencarian (wajib)', 'Gabungan Cabor + Nomor + Gender + Tipe'],
        ['Lokasi (Venue)', 'Dropdown pencarian (wajib)', 'Venue yang sudah terdaftar'],
        ['Waktu Pertandingan', 'Tanggal & Jam (wajib)', 'Jadwal mulai pertandingan'],
        ['Babak', 'Dropdown', 'Penyisihan / Perempat Final / Semifinal / Final'],
        ['Status', 'Dropdown', 'Terjadwal / Berlangsung / Selesai / Ditunda'],
    ])

    h3(doc, 'Susunan Peserta Resmi (Kunci Integritas Pertandingan)')
    para(doc, 'Sistem mewajibkan tepat 2 peserta (A vs B) dengan tipe yang seragam:')
    table(doc, ['Isian Peserta', 'Jenis Input', 'Aturan Penting'], [
        ['Jenis Peserta', 'Dropdown (seragam untuk A & B)', 'Kontingen / Individu / Tim — harus sama'],
        ['Peserta A — Kontingen', 'Dropdown pencarian (wajib)', 'Daerah afiliasi slot 1'],
        ['Peserta A — Nama Atlet', 'Teks (maks 100)', 'Wajib jika jenis = Individu'],
        ['Peserta A — Nama Tim', 'Teks (maks 150)', 'Wajib jika jenis = Tim'],
        ['Peserta B — Kontingen', 'Dropdown pencarian (wajib)', 'Daerah afiliasi slot 2'],
        ['Peserta B — Nama Atlet', 'Teks (maks 100)', 'Wajib jika jenis = Individu'],
        ['Peserta B — Nama Tim', 'Teks (maks 150)', 'Wajib jika jenis = Tim'],
    ])
    para(doc, 'Validasi: Peserta A dan B tidak boleh identik. Pembaruan peserta otomatis menyimpan data lama (soft-delete) sebelum diganti data baru.')

    # --- TAHAP 3 ---
    h2(doc, 'TAHAP 3: Hari Pertandingan — LiveScore Real-Time')
    para(doc, 'Saat peluit dibunyikan, Koresponden di pinggir lapangan menjalankan aksi:')
    num(doc, 'Login ke Web Admin dengan akun berperan "koresponden".')
    num(doc, 'Membuka menu LiveScore Center.')
    num(doc, 'Memilih pertandingan yang sedang berlangsung dari dropdown.')
    num(doc, 'Mengisi Skor A dan Skor B (angka non-negatif).')
    num(doc, 'Memilih status laga: Belum Mulai / Berlangsung / Istirahat / Selesai / Official.')
    num(doc, 'Menekan tombol "Simpan". Skor langsung terpancar ke seluruh layar publik tanpa jeda.')
    para(doc, '')
    bold_para(doc, 'Fitur Koreksi Skor (Jika Terjadi Kesalahan Input):')
    nbullet(doc, 'Centang checkbox "Ini adalah koreksi skor".')
    nbullet(doc, 'Isi alasan koreksi (minimal 5 karakter).')
    nbullet(doc, 'Riwayat koreksi tercatat permanen (append-only) dan tidak bisa dihapus.')

    # --- TAHAP 4 ---
    h2(doc, 'TAHAP 4: Penentuan Juara — Workflow Medali 4 Tahap')
    para(doc, 'Proses pencatatan medali dirancang anti-manipulasi dengan pemisahan tugas yang ketat:')
    table(doc, ['Tahap', 'Status', 'Pelaku', 'Aksi yang Dilakukan'], [
        ['1. Pengajuan', 'PENDING', 'Koresponden', 'Mengisi kontingen penerima, jumlah Emas/Perak/Perunggu, URL bukti, catatan → klik "Ajukan"'],
        ['2. Verifikasi', 'VERIFIED', 'Verifikator', 'Memeriksa kebenaran data → klik "Verifikasi" (sah) atau "Tolak" + alasan min 5 karakter'],
        ['3. Pengesahan', 'OFFICIAL', 'Super Admin', 'Klik "Publikasikan" → medali masuk ke Klasemen Umum yang tampil di Web Publik'],
        ['4. Penolakan', 'REJECTED', 'Verifikator / Super Admin', 'Dapat dilakukan dari PENDING atau VERIFIED → harus disertai alasan'],
    ])
    para(doc, 'Setiap aktor (submitted_by, verified_by, rejected_by, published_by) tercatat terpisah dan tidak bisa diubah.')
    doc.add_page_break()

    # ============================================================
    #  BAB V — FITUR WEB PUBLIK
    # ============================================================
    h1(doc, 'BAB V — Fitur Web Publik: Pengalaman Masyarakat')
    section_intro(doc, '', 'Tanpa perlu login, siapa saja dapat mengakses informasi pertandingan dari ujung jari mereka.')
    doc.add_paragraph('')

    stat_table(doc, [
        ('8', 'Halaman\nUtama'),
        ('SSE', 'Skor\nReal-Time'),
        ('PWA', 'Bisa\nDi-Install'),
        ('SEO', 'Ramah\nGoogle'),
    ])

    h2(doc, '5.1 Beranda (Landing Page)')
    nbullet(doc, 'Hero Banner Dinamis: Judul besar + teks sorotan gradient + latar parallax + badge "7 - 23 November 2026".')
    nbullet(doc, 'Hitung Mundur: 4 kotak (Hari, Jam, Menit, Detik) menuju Opening Ceremony. Otomatis hilang saat acara dimulai.')
    nbullet(doc, 'Profil Maskot: Toca (Merah — Semangat Juang) dan Toci (Putih — Sportivitas).')
    nbullet(doc, '4 Pintasan Utama: LiveScore, Jadwal, Medali, Cabor.')
    nbullet(doc, 'Venue Showcase: 4 venue teratas dengan auto-refresh 30 detik.')
    nbullet(doc, 'City Guide Bento Grid: Landmark Depok + statistik (40+ Destinasi, 120+ Kuliner).')

    h2(doc, '5.2 Cabang Olahraga')
    nbullet(doc, 'Halaman Katalog: Grid kartu seluruh cabor dengan ikon, nama, kategori, deskripsi.')
    nbullet(doc, 'Halaman Detail: Banner, nomor pertandingan, jadwal terkait, info teknis, venue terkait.')
    nbullet(doc, 'SEO Dinamis: Metadata (title, description, OpenGraph) dihasilkan otomatis per cabor.')

    h2(doc, '5.3 Venue Pertandingan')
    nbullet(doc, 'Peta Interaktif: Layout 2 kolom — daftar venue + peta Leaflet dengan marker seluruh lokasi.')
    nbullet(doc, 'Halaman Detail: Foto hero, badge kesiapan, kapasitas, jadwal, tombol rute Google Maps.')
    nbullet(doc, '6 Rekomendasi Terdekat: Algoritma Haversine menghitung jarak dari koordinat venue ke destinasi City Guide.')

    h2(doc, '5.4 Jadwal Pertandingan')
    nbullet(doc, 'Pencarian: Search bar kata kunci + 4 dropdown filter (Tanggal, Cabor, Venue, Status).')
    nbullet(doc, 'Pengelompokan otomatis berdasarkan Cabang Olahraga.')

    h2(doc, '5.5 LiveScore Center')
    nbullet(doc, 'Koneksi Ganda: Snapshot awal (REST) + pembaruan langsung (Server-Sent Events).')
    nbullet(doc, 'Validasi: Skor hanya muncul jika Peserta A & B sudah terisi lengkap.')
    nbullet(doc, 'Efek Visual: Kartu pertandingan berkedip saat menerima skor baru.')
    nbullet(doc, 'Panel Status: Koneksi SSE, jumlah update, tipe transport.')

    h2(doc, '5.6 Klasemen Medali')
    nbullet(doc, '3 Kartu Podium (Emas, Perak, Perunggu) + Tabel Klasemen lengkap.')
    nbullet(doc, 'Pengurutan: Emas ↓ → Perak ↓ → Perunggu ↓ → Nama Kontingen A-Z.')
    nbullet(doc, 'Hanya medali berstatus OFFICIAL yang tampil.')

    h2(doc, '5.7 City Guide (Panduan Kota)')
    nbullet(doc, '9 Kategori: Semua, Coffee Shop, Wisata Kuliner, Menginap, Wisata Buatan, Wisata Situ, Mall, RS, Lainnya.')
    nbullet(doc, '12 item per halaman dengan pagination.')

    h2(doc, '5.8 Fitur SEO & PWA')
    nbullet(doc, 'Sitemap XML, Robots.txt, OpenGraph, Twitter Card.')
    nbullet(doc, 'PWA: Bisa di-install di Home Screen, Service Worker, precaching aset.')
    doc.add_page_break()

    # ============================================================
    #  BAB VI — FITUR WEB ADMIN
    # ============================================================
    h1(doc, 'BAB VI — Fitur Web Admin: Senjata Panitia')
    section_intro(doc, '', 'Dapur operasional yang didesain untuk kecepatan kerja, keterbacaan, dan jejak audit.')
    doc.add_paragraph('')

    h2(doc, '6.1 Dashboard Operasional')
    para(doc, 'Tampilan pertama saat login — ringkasan kondisi sistem secara real-time:')
    table(doc, ['Kartu Statistik', 'Sumber Data'], [
        ['Total Cabang Olahraga', 'Real-time dari Master Data Service'],
        ['Total Venue', 'Real-time dari Venue Service'],
        ['Total City Guide', 'Real-time dari Master Data Service'],
        ['Total Kontingen', 'Real-time dari Master Data Service'],
        ['Total Atlet Terdaftar', 'Placeholder (siap diisi saat fitur atlet aktif)'],
        ['Medali Didistribusikan', 'Placeholder'],
        ['Skor Masuk Hari Ini', 'Placeholder'],
        ['Insiden Sistem', 'Placeholder'],
    ])
    nbullet(doc, 'Tabel Pertandingan Berlangsung: Waktu, Cabor, Pertandingan, Status.')
    nbullet(doc, '10 Log Audit Terkini: Warna hijau (CREATE), kuning (UPDATE), merah (DELETE).')

    h2(doc, '6.2 Master Data (6 Sub-Modul dalam 1 Workspace)')
    para(doc, 'Seluruh data inti dikelola dalam workspace bertab:')
    table(doc, ['Tab', 'Kolom Tabel', 'Aksi Tersedia'], [
        ['Cabang Olahraga', 'Nama, Kategori, Medali, TD, Status', 'Edit, Arsipkan'],
        ['Nomor Pertandingan', 'Cabor, Nomor, Gender, Tipe', 'Edit, Arsipkan'],
        ['Venue Depok', 'Nama, Alamat, Kapasitas', 'Edit, Arsipkan'],
        ['Data Kontingen', 'Logo, Nama, Tipe Daerah', 'Edit, Arsipkan'],
        ['Jadwal Pertandingan', 'Waktu, Cabor/Nomor, Peserta A vs B, Venue, Babak, Status', 'Edit, Arsipkan'],
        ['Recycle Bin', 'Nama, Jenis, Waktu Arsip, Aktor & Alasan', 'Pulihkan'],
    ])
    para(doc, 'Semua tabel mendukung: Pencarian, Pengurutan kolom, Pagination (10/25/50/100 baris).')

    h2(doc, '6.3 Recycle Bin (Pemulihan Data)')
    nbullet(doc, '8 entitas didukung: Cabor, Nomor Tanding, Kontingen, City Guide, Media, Hero, Venue, Jadwal.')
    nbullet(doc, 'Setiap penghapusan meminta alasan dan mencatat siapa yang menghapus.')
    nbullet(doc, 'Tombol "Pulihkan" mengembalikan data beserta seluruh relasinya.')

    h2(doc, '6.4 Hero Utama (Manajemen Banner Publik)')
    nbullet(doc, 'Judul (maks 180 karakter), Teks Sorotan (substring judul untuk efek gradient), Isi (maks 1200 karakter).')
    nbullet(doc, 'Gambar latar wajib dari Media Library. Hanya 1 Hero aktif pada satu waktu.')

    h2(doc, '6.5 Media Library (Galeri Aset Gambar)')
    nbullet(doc, 'Upload: Maks 10 MB, format JPG/PNG/WebP, nama file acak 16 byte.')
    nbullet(doc, 'Aksi: Salin URL, Arsipkan (soft-delete). Digunakan di seluruh modul sebagai pemilih gambar.')

    h2(doc, '6.6 Audit Log (CCTV Digital Sistem)')
    nbullet(doc, 'Mencatat SETIAP perubahan data: siapa, kapan, dari IP mana, apa yang diubah.')
    nbullet(doc, 'Integritas: SHA-256 hash per entry, database trigger mencegah modifikasi/penghapusan.')
    nbullet(doc, 'Export CSV untuk laporan pertanggungjawaban.')

    h2(doc, '6.7 Manajemen Akun (Khusus Super Admin)')
    nbullet(doc, 'Tambah user: Username, Email, Nama, Role, Password → tersinkronisasi 2 arah dengan Keycloak.')
    nbullet(doc, 'Hapus user: Menonaktifkan di Keycloak + soft-delete di database lokal.')
    doc.add_page_break()

    # ============================================================
    #  BAB VII — KEAMANAN
    # ============================================================
    h1(doc, 'BAB VII — Keamanan & Ketahanan Sistem')
    section_intro(doc, '', 'Keamanan bukan fitur tambahan — ia adalah fondasi utama sistem.')
    doc.add_paragraph('')

    table(doc, ['Aspek Keamanan', 'Implementasi', 'Manfaat'], [
        ['SSL/TLS', 'Wildcard SSL depok.go.id, TLS 1.2/1.3, HTTP/2', 'Data terenkripsi, Gembok Hijau di browser'],
        ['HSTS', 'max-age 2 tahun, includeSubDomains, preload', 'Browser selalu menggunakan HTTPS'],
        ['Content Security Policy', 'CSP dinamis per-path, evaluasi bertingkat', 'Mencegah injeksi skrip jahat (XSS)'],
        ['Rate Limiting', 'API: 30 req/s (burst 60), Auth: 5 req/s (burst 20)', 'Mencegah serangan brute-force'],
        ['Non-Root Containers', 'Semua service: user porprov/nextjs (UID 1001)', 'Meminimalkan dampak jika container diretas'],
        ['Docker Hardening', 'read_only, cap_drop ALL, no-new-privileges', 'Container tidak bisa mengubah filesystem'],
        ['Immutable Audit', 'PostgreSQL trigger melarang UPDATE/DELETE', 'Log tidak bisa dimanipulasi'],
        ['Append-Only Score', 'Trigger pada tabel livescore_revisions', 'Riwayat skor tidak bisa dihapus'],
        ['Transactional Outbox', 'Event + data dalam 1 transaksi DB', 'Event bisnis tidak pernah hilang'],
        ['Error Masking', 'Gateway menyaring error 5xx downstream', 'Detail implementasi tidak bocor ke browser'],
        ['Soft-Delete Mutlak', 'Semua data: deleted_at, deleted_by, delete_reason', 'Tidak ada penghapusan permanen akibat salah klik'],
        ['CORS Terpusat', 'Hanya Gateway yang mengelola CORS', 'Mencegah akses dari domain asing'],
    ])
    doc.add_page_break()

    # ============================================================
    #  BAB VIII — RBAC
    # ============================================================
    h1(doc, 'BAB VIII — Pemetaan Hak Akses (RBAC)')
    section_intro(doc, '', 'Setiap peran hanya melihat apa yang perlu dilihat, dan hanya bisa melakukan apa yang boleh dilakukan.')
    doc.add_paragraph('')

    table(doc, ['Menu / Fitur', 'Super Admin', 'Auditor', 'Admin Venue', 'Koresponden', 'Verifikator'], [
        ['Dashboard', '✅', '✅', '✅', '✅', '✅'],
        ['Master Data (CRUD)', '✅', '✅', '✅', '✅', '✅'],
        ['Hero Utama', '✅', '✅', '✅', '✅', '✅'],
        ['City Guide', '✅', '✅', '✅', '✅', '✅'],
        ['Media Library', '✅', '✅', '✅', '✅', '✅'],
        ['Profil Akun', '✅', '✅', '✅', '✅', '✅'],
        ['LiveScore Center', '✅', '—', '—', '✅', '—'],
        ['Pengajuan Medali', '✅', '—', '—', '✅', '—'],
        ['Verifikasi Medali', '✅', '—', '—', '—', '✅'],
        ['Publikasi Medali (OFFICIAL)', '✅', '—', '—', '—', '—'],
        ['Audit Log & Export CSV', '✅', '✅', '—', '—', '—'],
        ['Manajemen Akun Pengguna', '✅', '—', '—', '—', '—'],
        ['SSE Stream Admin', '✅', '✅', '—', '✅', '✅'],
    ])
    doc.add_page_break()

    # ============================================================
    #  BAB IX — INFRASTRUKTUR
    # ============================================================
    h1(doc, 'BAB IX — Infrastruktur & Deployment')
    section_intro(doc, '', '26 kontainer Docker berjalan bersama di atas satu server, diatur oleh satu perintah.')
    doc.add_paragraph('')

    h2(doc, '9.1 Komposisi Docker (26 Kontainer)')
    table(doc, ['Kategori', 'Jumlah', 'Rincian'], [
        ['Database & Cache', '3', 'PostgreSQL (7 DB), Redis, NATS JetStream'],
        ['Migrasi Database', '7', 'Migrasi otomatis saat startup (selesai lalu berhenti)'],
        ['Layanan Mikro Golang', '9', 'API Gateway, Master Data, Schedule, Venue, LiveScore, Medal, Audit, Realtime, User'],
        ['Aplikasi Web', '2', 'Public Web (Next.js), Admin Web (React/Nginx)'],
        ['Identitas & Bootstrap', '2', 'Keycloak + Bootstrap (realm, client, role, user awal)'],
        ['Reverse Proxy', '1', 'Nginx (SSL, routing, security headers, rate limiting)'],
        ['Observabilitas', '2', 'Prometheus (metrics) + Grafana (dashboard)'],
    ])

    h2(doc, '9.2 Keycloak SSO (Identitas Terpusat)')
    table(doc, ['Komponen', 'Detail'], [
        ['Realm', 'porprov'],
        ['Realm Roles (5)', 'super_admin, admin_venue, koresponden, verifikator, auditor'],
        ['Client: Admin Web', 'Public client, Authorization Code + PKCE S256'],
        ['Client: Mobile Admin', 'Public client, Standard Flow + Direct Access Grants'],
        ['Client: Backend Service', 'Confidential, Service Account (manage-users, view-users)'],
        ['User Awal', 'admin_depok (super_admin), koresponden_1 (koresponden)'],
    ])
    doc.add_page_break()

    # ============================================================
    #  BAB X — REKOMENDASI
    # ============================================================
    h1(doc, 'BAB X — Rekomendasi Pengembangan Selanjutnya')
    section_intro(doc, '', 'Ide-ide untuk membawa portal ini ke level berikutnya.')
    doc.add_paragraph('')

    table(doc, ['No', 'Fitur Rekomendasi', 'Manfaat yang Diharapkan', 'Prioritas'], [
        ['1', 'Aplikasi Mobile (React Native)', 'Pengalaman native Android/iOS, fitur kamera untuk scan QR atlet/tiket', 'Tinggi'],
        ['2', 'Push Notification & WhatsApp Bot', 'Notifikasi otomatis ke panitia (pengajuan medali, perubahan jadwal)', 'Tinggi'],
        ['3', 'Peta Venue Interaktif Terpadu', 'Satu peta digital seluruh venue + kepadatan + cuaca lokasi', 'Sedang'],
        ['4', 'Sponsor & Partnership Corner', 'Spot iklan dinamis di Web Publik, perputaran dari Admin', 'Sedang'],
        ['5', 'Dashboard Analitik Lanjutan', 'Grafik tren skor, statistik atlet, heat map pertandingan', 'Sedang'],
        ['6', 'Sistem Tiket Digital', 'Penjualan dan validasi tiket penonton secara elektronik', 'Rendah'],
    ])
    doc.add_page_break()

    # ============================================================
    #  BAB XI — KEBUTUHAN SISTEM
    # ============================================================
    h1(doc, 'BAB XI — Kebutuhan Sistem & Spesifikasi')

    table(doc, ['Komponen', 'Spesifikasi Minimal', 'Spesifikasi Rekomendasi'], [
        ['Sistem Operasi', 'Ubuntu 22.04 LTS', 'Ubuntu 24.04 LTS'],
        ['RAM', '8 GB', '16 GB'],
        ['CPU', '4 vCPU', '8 vCPU'],
        ['Penyimpanan', '50 GB SSD', '100 GB NVMe SSD'],
        ['Docker', 'Docker Engine 24+', 'Docker Engine 27+'],
        ['Domain', 'Subdomain aktif', 'porprov.depok.go.id (sudah aktif)'],
        ['SSL', "Let's Encrypt", 'Wildcard SSL depok.go.id (sudah terpasang)'],
        ['Bandwidth', '10 Mbps', '100 Mbps (untuk hari pertandingan)'],
    ])

    doc.add_paragraph('')
    doc.add_paragraph('')
    divider(doc)
    doc.add_paragraph('')
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_end.add_run('Dokumen ini disusun berdasarkan pembedahan langsung seluruh kode sumber repository porprov-depok, konfigurasi infrastruktur Docker, skema database, dan arsitektur layanan mikro yang telah berjalan di VPS Publik https://porprov.depok.go.id.')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Seluruh informasi akurat per Agustus 2026.')
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    out = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Laporan_Porprov_XV_2026.docx",
    )
    doc.save(out)
    print(f"DOCX saved: {out}")

    # Convert to PDF
    try:
        from docx2pdf import convert
        pdf = out.replace('.docx', '.pdf')
        convert(out, pdf)
        print(f"PDF saved: {pdf}")
    except Exception as e:
        print(f"PDF conversion note: {e}")

if __name__ == "__main__":
    main()
